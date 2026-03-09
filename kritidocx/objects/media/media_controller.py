"""
MEDIA CONTROLLER (The Asset Orchestrator)
-----------------------------------------
Responsibility:
Central entry point for creating Visual Elements (Images, Textboxes, Shapes).

Key Logic Flow:
1. Load Resource (ImageLoader)
2. Calculate Layout Math (PositioningEngine)
3. Calculate Styles (ShapeFactory)
4. Register with Document Part (generate rId)
5. Delegate Construction (XmlBuilder)

Features:
- Handles Floating vs Inline images logic automatically.
- Integrates CSS Sizing rules (%, px) with Word EMUs.
- Returns layout hooks for nested content (Textboxes).
"""

from kritidocx.xml_factory.drawing_xml import DrawingXml
from .image_loader import ImageLoader
from .positioning_engine import PositioningEngine
from .shape_factory import ShapeFactory
from kritidocx.xml_factory.xml_builder import XmlBuilder
from kritidocx.config.settings import AppConfig
import random

# Logging fallback
try:
    from kritidocx.utils.logger import logger
except ImportError:
    import logging
    logger = logging.getLogger("MyDocX_Media")

class MediaController:
    """
    Manages insertion of non-text visual objects.
    """

    def __init__(self, doc_driver_instance):
        """
        :param doc_driver_instance: Instance of DocxDriver (holding self.doc)
        """
        self.doc = doc_driver_instance.doc

    def add_image(self, src, container=None, style_data=None, alt_text=None):
        """
        Orchestrates adding an Image to the document.
        Updated: Resolves container early to allow fallback text injection on failure.
        """
        if getattr(AppConfig, 'DEBUG_MEDIA', False):
            logger.debug(f"\n🖼️ [MEDIA DEBUG] Processing Image: {src[:40]}...")
            logger.debug(f"   ➤ Raw CSS Received: {style_data}")
            # Border Specific check
            if style_data:
                logger.debug(f"   ➤ Border Check: border='{style_data.get('border')}', border-left='{style_data.get('border-left')}'")

      
        if not src: return
        if style_data is None: style_data = {}


        # 🔍 DEBUG LOG START
        debug_on = getattr(AppConfig, 'DEBUG_MEDIA', False)
        if debug_on:
            logger.debug(f"\n🖼️ [MEDIA START] Source: {src[:30]}... | Container: {type(container)}")

        # -------------------------------------------------------------
        # STEP 1: RESOLVE CONTAINER (Safety First)
        # -------------------------------------------------------------
        # We define where the image/text will go BEFORE trying to download.
        target_paragraph = container
        created_new_para = False

        # If passed Doc/Cell (Block), create a Paragraph.
        # If passed existing Paragraph, use it.
        if hasattr(container, 'add_paragraph') and not hasattr(container, 'add_run'):
            target_paragraph = container.add_paragraph()
            created_new_para = True
        
        # Fallback Safety: Ensure we have a valid paragraph object
        if not target_paragraph:
            try:
                 target_paragraph = self.doc.add_paragraph()
                 created_new_para = True
            except:
                 if debug_on: print("   ❌ Fatal: No Container found.")
                 return


        # -------------------------------------------------------------
        # STEP 2: LOAD & VALIDATE IMAGE
        # -------------------------------------------------------------
        # Gets local path + physical metadata
        image_path, meta = ImageLoader.get_processed_image(src,style_data)
        
        if debug_on:
            status = "✅ Found" if image_path else "❌ Failed"
            logger.debug(f"   📂 Path Resolution: {status} -> {image_path}")

        
        if not image_path:
            logger.error(f"Image Load Failed completely for: {src}")
            
            # [FIXED LOGIC] Inject Text Fallback instead of silent fail
            fallback_text = f"[IMG: {alt_text or 'Error'}]"
            run = target_paragraph.add_run(fallback_text)
            run.italic = True
            run.font.color.rgb = None # Optional: Set to Red using run manager if available
            return

        # -------------------------------------------------------------
        # STEP 3: CALCULATE DIMENSIONS (EMUs)
        # -------------------------------------------------------------
        native_w = meta.get('width', 100)
        native_h = meta.get('height', 100)

        target_width = PositioningEngine.calculate_width_emu(
            style_data, native_px_width=native_w
        )
        
        target_height = PositioningEngine.calculate_height_emu(
            style_data, 
            native_px_height=native_h,
            current_width_emu=target_width,
            native_px_width=native_w
        )

        # -------------------------------------------------------------
        # STEP 4: CALCULATE POSITIONING (Float/Absolute)
        # -------------------------------------------------------------
        pos_config = PositioningEngine.resolve_positioning(style_data)
        
        # Floating Object Logic: Collapse the host paragraph to avoid whitespace gaps
        if pos_config['is_floating'] and created_new_para:
            try:
                from docx.shared import Pt
                from docx.enum.text import WD_LINE_SPACING
                fmt = target_paragraph.paragraph_format
                fmt.space_after = Pt(0)
                fmt.space_before = Pt(0)
                fmt.line_spacing = Pt(0)
                fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            except Exception: pass

        # -------------------------------------------------------------
        # STEP 5: EXECUTE INSERTION
        # -------------------------------------------------------------
        try:
            
            if debug_on: print(f"   🚀 Sending to XMLBuilder...")

            
            img_id_ref = "Img_" + str(abs(hash(src)))[:8]
            
            XmlBuilder.insert_image(
                container=target_paragraph,
                image_path=image_path,
                width_emu=target_width,
                height_emu=target_height,
                pos_config=pos_config,
                image_id=img_id_ref,
                style_data=style_data
            )

            # ✅ [NUCLEAR FIX]: DIRECT XML OVERRIDE FOR SPACING
            # Python wrappers पर भरोसा करने के बजाय सीधे XML लिखें।
            # 'line_rule="auto"' और 'line="240"' (100%) इमेज को सांस लेने की जगह देता है।
            
            if not pos_config.get('is_floating'):
                # Spacing के लिए सीधे Builder को कॉल करें
                # यह पुराने किसी भी "Exact/20twips" टैग को ओवरराइट कर देगा।
                XmlBuilder.set_paragraph_spacing(
                    target_paragraph, 
                    line_rule='auto',  # Height expands automatically
                    line=240,          # Standard Single Line
                    before=0,          
                    after=100          # हल्का सा बॉटम स्पेस (5pt) ताकि इमेज नीचे चिपके नहीं
                )



            if debug_on: print(f"   ✨ XML Injection Complete.")

        except Exception as e:
            if debug_on:
                import traceback
                logger.error(f"   🔥 [MEDIA CRASH]: {e}")
                print(traceback.format_exc()) # असली गुनहगार यह बताएगा
                
            logger.error(f"Image Insertion Logic Failed: {e}")
            target_paragraph.add_run(f"[IMG ERROR]").italic = True

    def add_floating_textbox(self, node_info, container=None):
        """
        [ENGINE CORE]: Creates an Anchored Textbox (Shape).
        Flow: CSS -> Positioning -> Dimensions -> XML Factory -> Ghost Paragraph.
        
        Returns:
            The INTERNAL content node (w:txbxContent) allows the Router 
            to fill it recursively with text/tables.
        """
        style_data = node_info.get('style_dict', {})
        
        # 1. LAYOUT MATH (Step 3 Engine)
        # यह तय करेगा कि absolute है या relative, और wrap कैसा होगा
        pos_config = PositioningEngine.resolve_positioning(style_data)
        
        # Guard Clause: अगर फ्लोटिंग नहीं है और न ही रोटेशन है, तो यह साधारण DIV है
        # इसे राउटर को सामान्य तरह से प्रोसेस करने दें
        if not pos_config['is_floating'] and pos_config['rotation'] == 0:
            return None # Signal to Router to treat as normal Block

        # 2. DIMENSIONS (Fallback Logic)
        # Textboxes need strict EMUs. 
        # Default: 2.5 inch Width, 1.25 inch Height (Word standard fallback)
        def_w_emu = 2286000 
        def_h_emu = 1143000
        
        # CSS या Px को EMUs में बदलें
        # नोट: हम PositioningEngine का उपयोग करके 'auto' को handle कर रहे हैं
        w_emu = PositioningEngine.calculate_width_emu(style_data, 200) or def_w_emu
        
        # Height is tricky. HTML 'auto' means 'grow', but Shape needs base size.
        h_emu = PositioningEngine.calculate_height_emu(style_data, 100, w_emu, 200) or def_h_emu
        
        # config में आयाम सेट करें (XML Anchor के लिए)
        pos_config['width'] = w_emu
        pos_config['height'] = h_emu

        # -------------------------------------------------------------
        # ✅ [CRITICAL LAYOUT FIX]: Right/Bottom Anchor Math
        # -------------------------------------------------------------
        # Word एंकर को Top-Left से मापता है।
        # यदि हम 'Right Margin' के सापेक्ष हैं, तो हमें (Offset + Width) घटाना होगा
        # ताकि बॉक्स का Right Edge उस ऑफसेट पर आए, न कि Left Edge.
        
        # 1. Horizontal Correction (Right: 50px -> Move Left by (50 + Width))
        if pos_config.get('rel_h') == 'rightMargin' and pos_config.get('pos_x', 0) < 0:
            # pos_x पहले से Negative (-50px) है, इसमें Width भी घटा दें
            pos_config['pos_x'] = pos_config['pos_x'] - w_emu

        # 2. Vertical Correction (Bottom: 50px -> Move Up by (50 + Height))
        if pos_config.get('rel_v') == 'bottomMargin' and pos_config.get('pos_y', 0) < 0:
            # pos_y पहले से Negative (-50px) है, इसमें Height भी घटा दें
            pos_config['pos_y'] = pos_config['pos_y'] - h_emu
            
        # -------------------------------------------------------------

        # 3. VISUAL STYLE (Step 2 Engine)
        # Background, Border, Padding Calculation
        shape_style = ShapeFactory.create_shape_config(style_data, shape_type='rect')

        # 4. IDENTITY
        # Unique Integer ID (Important for Z-Ordering consistency)
        # हम simplePos में random use कर रहे हैं, ड्राइंग फैक्ट्री अपना इंटरनल काउंटर संभालेगी
        uid = random.randint(10000, 999999)

        # 5. CONTAINER STRATEGY: THE GHOST PARAGRAPH
        target_paragraph = None
        
        # क्या हमें नया पैराग्राफ बनाना चाहिए?
        if hasattr(container, 'add_paragraph'):
            target_paragraph = container.add_paragraph()
            
            # --- 🔥 THE GHOST TRICK ---
            # पैराग्राफ का वजूद मिटा दें ताकि पेज पर जगह न ले (Height = 0)
            if pos_config['is_floating']:
                from docx.shared import Pt
                from docx.enum.text import WD_LINE_SPACING
                
                pf = target_paragraph.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                # 'Exactly 0' forces the line height to nothing
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(0) 
        else:
            # Fallback (Existing paragraph)
            target_paragraph = container

        # 6. XML CONSTRUCTION (Step 1 Factory Integration)
        try:
            # [SMART AUTOFIT FIX]
            # चेक करें कि CSS में हाइट फिक्स है या नहीं
            css_h = style_data.get('height')
            # अगर हाइट 'auto' है या दी ही नहीं गई है, तो AutoFit (Resize) करें
            is_autofit = (css_h is None or str(css_h) == 'auto')

            # A. Graphic Body (WSP Element)
            graphic_xml, content_xml = DrawingXml.create_textbox_structure(
                unique_id=uid,
                cx=w_emu,
                cy=h_emu,
                fill_color=shape_style['fill'].get('color') if shape_style['fill'].get('type') == 'solid' else None,
                border_color=shape_style['outline'].get('color') if shape_style['outline'].get('type') == 'solid' else None,
                dash_style=shape_style['outline'].get('dash', 'solid'),
                fit_to_text=is_autofit, # ✅ यहाँ डायनामिक वैल्यू पास करें
                rotation=pos_config['rotation'],
                shadow_config=shape_style.get('shadow')
            )
            
            # ---------------------------------------------------------
            # B. Anchor Wrapper (WP Element) - FIXED SECTION
            # ---------------------------------------------------------
            
            # 1. आयामों (Dimensions) का डिक्शनरी बनाएं
            dims_dict = {'width_emu': w_emu, 'height_emu': h_emu}

            # 2. यह ग्राफ़िक को पेज पर सही जगह चिपकाएगा (Step 3 pos_config के अनुसार)
            # wrap_type 'none' means Overlay (CSS behavior) vs 'square' (Text push)
            use_square_wrap = (pos_config.get('wrap_type') == 'square')
            
            anchor_xml = DrawingXml.create_absolute_anchor(
                graphic_element=graphic_xml,
                extent_info=dims_dict,        # <--- 🔥 ADDED THIS MISSING ARGUMENT
                coords=pos_config,
                unique_id=uid,
                force_wrap=use_square_wrap
            )
            
            # 7. INJECTION (Manual)
            # Textbox के लिए image relationship की जरूरत नहीं है, सीधे XML इंजेक्ट करें।
            
            # --- Manual Inject Fix (Safe & Direct) ---
            run = target_paragraph.add_run()
            
            # Collapse run text size (Extra safety for ghost paragraph)
            from docx.shared import Pt
            run.font.size = Pt(0.5) 
            
            # <w:drawing> parent बनाना
            drawing_node = XmlBuilder.create_element('w:drawing')
            drawing_node.append(anchor_xml)
            
            # Run के अंदर डालें (Using low-level append for direct XML)
            if hasattr(run, '_element'):
                run._element.append(drawing_node)
            else:
                # Fallback if accessed differently
                run._r.append(drawing_node)

            return content_xml # Success: Return the pocket for Router to fill!
        
        except Exception as e:
            if getattr(AppConfig, 'DEBUG_MEDIA', False):
                import traceback
                logger.error(f"   🔥 [TBX ERROR]: {e}")
                logger.debug(traceback.format_exc())
            return None