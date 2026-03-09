# ------------------------------------------------------------
# UPDATE IMPORTS (Top of the file)
# ------------------------------------------------------------
from kritidocx.utils.logger import logger
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.shared import Pt, RGBColor

# नए पावरफुल मैनेजर्स को शामिल करें
from kritidocx.config.settings import AppConfig
from kritidocx.objects.text.run_manager import RunManager
from kritidocx.objects.text.paragraph_manager import ParagraphManager
from kritidocx.basics.unit_converter import UnitConverter
from kritidocx.xml_factory.xml_builder import XmlBuilder

class HeaderFooterManager:
    """
    Manages Headers, Footers, and Page Numbers.
    Supports: Different First Page, Odd/Even Pages.
    """

    @staticmethod
    def configure(section, different_first_page=False, different_odd_even=False):
        section.different_first_page_header_footer = different_first_page
        section.even_and_odd_headers = different_odd_even

    # ------------------------------------------------------------
    # REPLACE FUNCTION: add_content (Inside class HeaderFooterManager)
    # ------------------------------------------------------------
    @staticmethod
    def add_content(section, content_type='header', is_first_page=False, text=None, image_path=None, style_data=None):
        """
        Generic method using the full power of RunManager and ParagraphManager.
        Supports HTML CSS fully (Colors, Borders, Alignment, Fonts).
        """
        if style_data is None: style_data = {}

        if getattr(AppConfig, 'DEBUG', False):
            logger.debug(f"   🎨 [HF MANAGER] Applying {content_type.upper()} Styling...")
            logger.debug(f"      -> Received Color: {style_data.get('color')}")
            logger.debug(f"      -> Received Align: {style_data.get('text-align')}")



        target = None
        if content_type == 'header':
            target = section.first_page_header if is_first_page else section.header
            # डिफॉल्ट (अगर HTML में text-align न हो)
            if 'text-align' not in style_data: style_data['text-align'] = 'right'
        else:
            target = section.first_page_footer if is_first_page else section.footer
            if 'text-align' not in style_data: style_data['text-align'] = 'center'

        # पुराना कंटेंट साफ़ करें
        if len(target.paragraphs) > 0:
            p = target.paragraphs[0]
            p.text = "" 
        else:
            p = target.add_paragraph()

        # 1. PARAGRAPH STYLE (Borders & Alignment)
        # ParagraphManager को बॉर्डर (underline) और अलाइनमेंट संभालने दें
        ParagraphManager.apply_formatting(p, style_data)

        # 2. TEXT STYLE (Color, Font, Size, Bold)
        if text:
            # RunManager अब अपने आप ColorManager और FontHandler का उपयोग करेगा
            RunManager.create_run(p, text, style_data)

        # 3. IMAGE HANDLING
        if image_path:
            img_run = p.add_run()
            try:
                # 1500000 EMUs = approx 1.5 inches width standard logo
                img_run.add_picture(image_path, width=1500000) 
            except Exception as e:
                pass # Silent fail logs internally elsewhere

        return p
    
    # ------------------------------------------------------------
    # REPLACE FUNCTION: add_page_numbers
    # ------------------------------------------------------------
    @staticmethod
    def add_page_numbers(section, style_data=None):
        """
        Appends page number with consistent styling via RunManager.
        """
        footer = section.footer
        
        # डिफ़ॉल्ट ग्रे स्टाइल अगर कोई CSS नहीं मिली
        fallback_style = {'color': 'gray', 'font-size': '9pt'}
        final_style = style_data if style_data else fallback_style

        # सही पैराग्राफ ढूँढें
        if len(footer.paragraphs) > 0:
            p = footer.paragraphs[0]
            
            # Separator " | " उसी स्टाइल (रंग/फॉन्ट) में जोड़ें
            RunManager.create_run(p, " | ", final_style)
            
        else:
            p = footer.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Field Code को डायरेक्ट XML Builder के ज़रिए जोड़ें
        # नोट: Field Code ({PAGE}) रन के स्टाइल को Inherit करता है,
        # लेकिन रंग सुनिश्चित करने के लिए हम XML स्तर पर run बना सकते हैं,
        # अभी के लिए साधारण Insertion ही काफ़ी है क्योंकि यह पिछले रन के स्टाइल को फॉलो करता है।
        
        # लेकिन बेहतर दिखने के लिए एक Empty formatted run जोड़ें
        # जिससे PAGE नंबर उसी फॉर्मेट में आए।
        phantom_run = RunManager.create_run(p, "", final_style)
        
        # फील्ड कोड डालें
        XmlBuilder.insert_field_code(p, "PAGE")
        
        
    @staticmethod
    def get_active_header(section, is_first_page=False):
        """
        Returns the raw Header object (cleared of default empty paragraphs).
        Allows Router to insert Tables/Images directly.
        """
        # 1. Target चुनें
        target = section.first_page_header if is_first_page else section.header
        
        # 2. सफाई अभियान (Cleanup)
        # Word डिफ़ॉल्ट रूप से एक खाली पैराग्राफ देता है। अगर हम अपना HTML Table डालेंगे,
        # तो वह इस पैराग्राफ के नीचे आएगा (एक्स्ट्रा स्पेस)। इसे हटाना जरुरी है।
        for p in target.paragraphs:
            p._element.getparent().remove(p._element)
            
        return target

    # [ADD THIS NEW METHOD FOR ROUTER SUPPORT]
    @staticmethod
    def get_active_footer(section, is_first_page=False):
        """Same logic for Footer."""
        target = section.first_page_footer if is_first_page else section.footer
        
        # Cleanup
        for p in target.paragraphs:
            p._element.getparent().remove(p._element)
            
        return target

    # [ADD HELPER FOR PAGE NUMBER INJECTION]
    @staticmethod
    def add_page_numbers_to_container(container, style_data=None):
        """Used when we want auto-numbering even with rich content."""
        p = container.add_paragraph()
        p.alignment = 1 # Center (import enum if needed)
        
        # Apply CSS color if present
        if style_data and style_data.get('color'):
            from kritidocx.objects.text.run_manager import RunManager
            RunManager.create_run(p, "", style_data) # Phantom run to set style
            
        from kritidocx.xml_factory.xml_builder import XmlBuilder
        XmlBuilder.insert_field_code(p, "PAGE")