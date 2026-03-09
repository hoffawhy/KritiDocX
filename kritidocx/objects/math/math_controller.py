"""
MATH CONTROLLER MODULE (The Equation Orchestrator)
--------------------------------------------------
Responsibility:
Central entry point for inserting Equations.
Decides whether to process as LaTeX or MathML and handles document injection.

Process Flow:
1. Validate Input (LaTeX String or BS4 Tag).
2. Clean Input (via LatexParser).
3. Convert to OMML (via OmmlEngine).
4. Apply Styles (Color/Fonts via StyleApplicator).
5. Inject into Document (Inline vs Block positioning).

Key Feature:
Fail-Safe Mechanism: If conversion fails, it gracefully degrades to plain italic text.
"""

import logging
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .omml_engine import OmmlEngine
from .latex_parser import LatexParser
from .style_applicator import StyleApplicator
from kritidocx.xml_factory.xml_builder import XmlBuilder

logger = logging.getLogger("MyDocX_Math")

class MathController:
    """
    Main Logic Class for Mathematics.
    """

    def __init__(self, doc_driver):
        self.doc = doc_driver.doc
        # Instantiate engine (loads XSLT once)
        self.engine = OmmlEngine()

    def process_math(self, content, container, style_data=None, is_latex=True):
        """
        Orchestrates equation processing.

        Args:
            content: Raw string (LaTeX) OR BeautifulSoup Tag (MathML/Span).
            container: Paragraph (Inline) or Document/Cell (Block).
            style_data: CSS dict for colors/fonts.
            is_latex (bool): True for $$..$$, False for <math>.
        """
        if not content: return

        # 1. Fallback Check: Is the engine healthy?
        if not self.engine._is_ready:
            raw_text = str(content.get_text()) if hasattr(content, 'get_text') else str(content)
            self._insert_fallback_text(raw_text, container, reason="Missing Dependencies")
            return

        try:
            # 2. Conversion Strategy
            omml_element = None

            if is_latex:
                # A. LaTeX Mode ($$ E=mc^2 $$)
                clean_tex = LatexParser.normalize(content)
                omml_element = self.engine.convert_to_omml(clean_tex)
            else:
                # B. MathML Mode (<math>...</math>)
                # [UPDATED LOGIC]
                
                # 1. Check Layout Attribute (display="block")
                # If content is a BS4 Tag, extract attributes
                if hasattr(content, 'get'):
                    display_mode = content.get('display')
                    if display_mode == 'block':
                        # Update style_data to force block alignment later
                        if style_data is None: style_data = {}
                        style_data['display'] = 'block'

                # 2. Sanitize & Convert
                clean_mml = self._sanitize_mathml_string(content)
                omml_element = self.engine.convert_to_omml(clean_mml)


            # Validation
            if omml_element is None:
                raise ValueError("XSLT returned empty OMML result")

            # 3. Styling (Inject Colors)
            if style_data:
                StyleApplicator.apply_style(omml_element, style_data)

            # 4. Injection & Layout Strategy
            # Is this Block Level ($$..$$) or Inline ($..$)?
            # Regex or attributes usually decide this before calling process_math, 
            # but we reinforce logic here based on Container type.

            self._inject_into_document(omml_element, container, style_data)

        except Exception as e:
            logger.warning(f"Math Render Failed: {e}")
            # Ensure text is readable at least
            raw_fallback = str(content)
            self._insert_fallback_text(raw_fallback, container, reason="Render Error")

    # -------------------------------------------------------------------------
    # 🧩 INTERNAL LOGIC
    # -------------------------------------------------------------------------

    def _inject_into_document(self, omml_node, container, style_data):
        """
        Smartly inserts the equation based on container type.
        """
        target_paragraph = None
        is_block_mode = False

        # --- A. DETECT CONTEXT ---
        if hasattr(container, 'add_run'):
            # It's an existing Paragraph (Inline Context)
            target_paragraph = container
        
        elif hasattr(container, 'add_paragraph'):
            # It's a Document/Cell/Div (Block Context)
            target_paragraph = container.add_paragraph()
            is_block_mode = True
            
        else:
            # Raw XML Element (Textbox inner content) fallback
            from docx.text.paragraph import Paragraph
            # Requires complex init, simpler to assume wrapper provided
            logger.error("Invalid Container for Math. Skipping.")
            return

        # --- B. BLOCK LAYOUT RULES ---
        # Check explicit CSS override
        css_display = style_data.get('display', '').lower() if style_data else ''
        if css_display == 'block' or css_display == 'inline-block':
             # Even if inline para, maybe we want to break line? 
             # For safety, standard math behavior is inline unless new para created.
             if is_block_mode:
                 target_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # --- C. XML INSERTION ---
        # Word needs Math to be part of the Paragraph content tree
        # Use XmlBuilder to append the raw OxmlElement properly.
        XmlBuilder.append_child(target_paragraph, omml_node)

    def _insert_fallback_text(self, text, container, reason="Error"):
        """
        Creates a visible text representation if Math fails.
        Format: [Equation: E=mc^2]
        """
        if hasattr(container, 'add_run'):
            p = container
        elif hasattr(container, 'add_paragraph'):
            p = container.add_paragraph()
        else:
            return 

        run = p.add_run(f" [Formula: {text}] ")
        run.italic = True
        run.font.color.rgb = None # Inherit default or make Red if critical
        
        logger.info(f"Inserted Fallback for Math ({reason})")
        
        
    def _sanitize_mathml_string(self, content_node):
        """
        [NEW HELPER] Ensures MathML has correct namespaces for XSLT.
        """
        # 1. Convert BS4 Tag to String
        xml_str = str(content_node)

        # 2. FORCE NAMESPACE (Namespace Injection)
        # MML2OMML.XSL requires the strict namespace to match templates.
        if 'xmlns="http://www.w3.org/1998/Math/MathML"' not in xml_str:
            # Inject namespace if missing
            xml_str = xml_str.replace('<math', '<math xmlns="http://www.w3.org/1998/Math/MathML"', 1)

        # 3. Handle 'mml:' prefix without definition
        if 'mml:' in xml_str and 'xmlns:mml' not in xml_str:
            xml_str = xml_str.replace('<math', '<math xmlns:mml="http://www.w3.org/1998/Math/MathML"', 1)

        return xml_str