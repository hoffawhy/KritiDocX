"""
HEADING MANAGER MODULE (The Structure Architect)
------------------------------------------------
Responsibility:
Creates structured Headings (H1-H9) for document navigation.

Key Features:
1. Structural Integrity: Enforces 'KeepWithNext' to prevent orphaned headings.
2. Outline Levels: Ensures headings appear in Word's Navigation Pane.
3. Bookmarking: Auto-generates anchors for TOC jumping.
4. Smart Overrides: Blends Word Native Styles with CSS Overrides.

Integration:
- Uses 'ParagraphManager' for block properties.
- Uses 'RunManager' for inner text logic.
"""

import re
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from kritidocx.xml_factory.xml_builder import XmlBuilder
from .paragraph_manager import ParagraphManager
from .run_manager import RunManager
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import random


class HeadingManager:
    """
    Controller for Outline Levels and Document Headings.
    """

    @classmethod
    def add_heading(cls, container, level, text_or_node, style_data=None, auto_bookmark=True):
        """
        Creates a Heading with logic.
        
        Args:
            container: Document or Cell (Note: Headings inside tables are allowed but quirky).
            level (int): 1 to 9.
            text_or_node: String content OR BeautifulSoup Tag.
            style_data (dict): CSS Properties overrides.
            auto_bookmark (bool): Create a bookmark anchor based on text?
        
        Returns:
            The created Paragraph object.
        """
        if not style_data: style_data = {}

        # 1. Level Validation
        # Word supports Heading 1-9
        safe_level = max(1, min(9, int(level)))
        style_name = f'Heading {safe_level}'

        # 2. Create Paragraph with Base Style
        # [CRITICAL FIX]: Support Raw XML Containers (Textboxes/Shapes)
        from docx.text.paragraph import Paragraph
        
        paragraph = None
        
        if hasattr(container, 'add_paragraph'):
            # CASE A: Standard Document/Cell Object
            paragraph = container.add_paragraph('', style=style_name)
            
        elif hasattr(container, 'append'):
            # CASE B: Raw XML Element (Textbox Content)
            # मैन्युअल रूप से पैराग्राफ नोड बनाएं
            p_node = OxmlElement('w:p')
            
            # स्टाइल सेट करने के लिए pPr जोड़ें
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), style_name)
            pPr.append(pStyle)
            p_node.append(pPr)
            
            # कंटेनर में जोड़ें
            container.append(p_node)
            
            # पाइथन रैपर बनाएं (taaki formatting apply ho sake)
            paragraph = Paragraph(p_node, None)
            
        else:
            # यदि कुछ समझ नहीं आया, तो चुपचाप वापस लौट जाएं (Safety)
            return None

        # [COLOR/ALIGN FIX]: स्टाइल अप्लाई करें
        if style_data.get('color') or style_data.get('text-align'):
            ParagraphManager.apply_formatting(paragraph, style_data)
            
            
        # 3. ENFORCE STRUCTURE (Critical for Professional Docs)
        cls._apply_structural_rules(paragraph, safe_level, style_data)

        # 4. PROCESS CONTENT & COLOR OVERRIDES
        # Headings often need specific color overrides (e.g. Branding colors) 
        # even if using a preset style.
        
        # Check text or object
        # Note: If it's a node, caller (Router) usually iterates children.
        # But here we act as the 'endpoint' logic for simpler H tags.
        raw_text = ""
        if hasattr(text_or_node, 'get_text'):
            raw_text = text_or_node.get_text().strip()
            # If style_data was passed from Router, it's already robust.
        else:
            raw_text = str(text_or_node).strip()

        if raw_text:
            run = RunManager.create_run(paragraph, raw_text, style_data)
            
            # Explicit Override: Native styles usually enforce a color. 
            # If User provided explicit color, RunManager applied it. 
            # If RunManager saw nothing, it used defaults.
            # No extra action needed here unless we want to force 'Heading 1' blue explicitly via Theme.

        # 5. BLOCK FORMATTING (Borders, Alignment, Spacing)
        ParagraphManager.apply_formatting(paragraph, style_data)

        # 6. BOOKMARKING (For TOC / Navigation) - [UPDATED FIX]
        if auto_bookmark and raw_text:
            # Slugify: "My Report 2025" -> "My_Report_2025"
            safe_slug = re.sub(r'[\W_]+', '_', raw_text).strip('_')
            unique_mark = f"{safe_slug}_{random.randint(100, 999)}"
            
            # Bookmark ID (Integer string)
            bm_id = str(random.randint(10000000, 99999999))
            
            # Create Start Tag: <w:bookmarkStart w:id="..." w:name="..."/>
            start = OxmlElement('w:bookmarkStart')
            start.set(qn('w:id'), bm_id)
            start.set(qn('w:name'), unique_mark)
            
            # Create End Tag: <w:bookmarkEnd w:id="..."/>
            end = OxmlElement('w:bookmarkEnd')
            end.set(qn('w:id'), bm_id)
            
            # Inject: Start at beginning, End at end of paragraph content
            paragraph._p.insert(0, start)
            paragraph._p.append(end)
        return paragraph

    @staticmethod
    def _apply_structural_rules(paragraph, level, style_data):
        """
        Apply logic that keeps document unbreakable.
        """
        p_fmt = paragraph.paragraph_format
        
        # Rule A: Keep With Next
        # Headings must always stick to the paragraph below them.
        p_fmt.keep_with_next = True
        
        # Rule B: Level 1 Start on New Page (Optional Config logic)
        # Check if style request asks for break OR standard behavior
        if level == 1 and not style_data.get('disable_page_break'):
            # Only apply if it's NOT the very first element of document (avoid blank page 1)
            # Accessing document element list is costly, we assume intelligent caller
            # Or explicit css 'page-break-before: always' handling in ParagraphManager covers this.
            pass 

        # Rule C: Outline Level
        # Usually style 'Heading X' handles this automatically. 
        # But if we inject custom XML, we re-enforce it.
        XmlBuilder.set_paragraph_outline_level(paragraph, level - 1) # Word logic: lvl 0 = H1

    # =========================================================================
    # ⚙️ EXTRA UTILITIES
    # =========================================================================

    @staticmethod
    def create_toc_placeholder(doc):
        """
        Generates the standard Word TOC field code.
        Placement usually done by Router via <div id="toc">.
        """
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Call Factory to build { TOC \o "1-3" \h \z \u }
        XmlBuilder.insert_field_code(p, 'TOC \\o "1-3" \\h \\z \\u')
        return p