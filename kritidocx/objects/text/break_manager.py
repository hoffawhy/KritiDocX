"""
from kritidocx.utils.logger import logger
BREAK MANAGER MODULE (The Flow Controller)
------------------------------------------
Responsibility:
Manages Line Breaks, Page Breaks, and Column Breaks.

Crucial Feature: "Clearance Logic"
Just like CSS `clear: both`, Word has <w:br w:clear="all"/>.
This module implements that to fix Layout issues around floating images.

Types:
1. Line Break (Shift+Enter)
2. Page Break (Ctrl+Enter)
3. Column Break (Next Column)
4. Text Wrapping Break (With Clear option)
"""

from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.run import Run
from docx.text.paragraph import Paragraph

from kritidocx.utils import logger

class BreakManager:
    """
    Controller for Document Flow Interruptions.
    """

    # CSS 'clear' values map to Word XML attribute values
    CLEAR_MAP = {
        'left': 'left',
        'right': 'right',
        'both': 'all',
        'all': 'all'
    }

    # Standard Break Types Enum Map
    BREAK_TYPES = {
        'line': WD_BREAK.LINE,
        'page': WD_BREAK.PAGE,
        'column': WD_BREAK.COLUMN,
        'section_next': WD_BREAK.SECTION_NEXT_PAGE, # Less common via run, handled by Layout usually
        'wrap': WD_BREAK.TEXT_WRAPPING # Special type
    }

    @staticmethod
    def add_break(container, break_type='line', count=1, clear_mode=None):
        """
        Master method to insert breaks.
        
        Args:
            container: The Document, Paragraph, or Run to attach to.
            break_type (str): 'line', 'page', 'column', 'wrap'.
            count (int): How many times to repeat (e.g., <br><br>).
            clear_mode (str): 'left', 'right', 'all' (equivalent to CSS clear:both).
        """
        if count < 1: return

        # 1. Resolve Target Run
        # We need a 'Run' object to insert a break.
        # If user passed Document/Paragraph, find/create the run.
        target_run = BreakManager._resolve_run(container)
        
        if not target_run:
            logger.error("⚠️ BreakManager Error: Could not resolve a valid Text Run target.")
            return

        # 2. Advanced Feature: "Clear" Break (CSS clear: both logic)
        # Standard python-docx add_break() doesn't support 'clear' attributes nicely.
        # We must manipulate XML manually if 'clear_mode' is present.
        
        if clear_mode and break_type in ['line', 'wrap']:
            xml_clear_val = BreakManager.CLEAR_MAP.get(clear_mode.lower())
            
            if xml_clear_val:
                for _ in range(count):
                    # Manual XML Injection
                    br = OxmlElement('w:br')
                    if break_type == 'wrap':
                        br.set(qn('w:type'), 'textWrapping')
                    
                    br.set(qn('w:clear'), xml_clear_val)
                    target_run._r.append(br)
                return

        # 3. Standard Break
        wd_break_type = BreakManager.BREAK_TYPES.get(break_type, WD_BREAK.LINE)
        
        for _ in range(count):
            target_run.add_break(wd_break_type)

    # -------------------------------------------------------------------------
    # 🕵️ INTERNAL HELPER
    # -------------------------------------------------------------------------

    @staticmethod
    def _resolve_run(container):
        """
        Intelligently finds where to put the break.
        Doc -> Add Para -> Add Run
        Para -> Add Run (or use last)
        Run -> Use Run
        """
        # Case A: Input is already a Run
        if isinstance(container, Run):
            return container

        # Case B: Input is a Paragraph
        if isinstance(container, Paragraph):
            # Optimisation: Use last run if exists, else create new
            if container.runs:
                return container.runs[-1]
            else:
                return container.add_run()

        # Case C: Input is Document (or Cell, Header, etc - anything with add_paragraph)
        if hasattr(container, 'add_paragraph'):
            para = container.add_paragraph()
            return para.add_run()

        return None

    # -------------------------------------------------------------------------
    # ⚡ QUICK ALIASES (Sugar Syntax)
    # -------------------------------------------------------------------------

    @classmethod
    def apply_page_break(cls, container):
        """Standard Page Break shortcut."""
        cls.add_break(container, 'page')

    @classmethod
    def apply_column_break(cls, container):
        """Moves text to next column (in multi-column layouts)."""
        cls.add_break(container, 'column')

    @classmethod
    def apply_clearing_break(cls, container):
        """
        [Special Helper] 
        Adds a break that pushes content below any floating images.
        Used after <img> tags or floating divs.
        """
        cls.add_break(container, 'line', clear_mode='all')