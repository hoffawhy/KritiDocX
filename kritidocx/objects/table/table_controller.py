"""
TABLE CONTROLLER (The Grid Orchestrator)
----------------------------------------
Responsibility:
Integrates all Table Sub-modules (Matrix, Props, Rows, Cells) into one pipeline.

Architecture (Flow of Control):
1. Router -> Calls TableController.process_table()
2. TableController -> Calls MatrixEngine (To calculate geometry)
3. TableController -> Creates Empty Word Table
4. TableController -> Calls PropsManager (For global width/borders)
5. TableController -> Iterates Rows & Cells:
    -> Calls RowManager (For height/headers)
    -> Calls CellManager (For shading/merging)
        -> CellManager Calls Router Back (Callback) to fill content.

Dependecies:
- Requires a valid 'router_instance' passed at runtime to handle recursive content.
"""

import logging
from kritidocx.config.settings import AppConfig
from kritidocx.objects.table.matrix_engine import MatrixEngine
from kritidocx.objects.table.props_manager import TablePropsManager
from kritidocx.objects.table.row_manager import RowManager
from kritidocx.objects.table.cell_manager import CellManager
from docx.shared import Inches, Twips
from docx.oxml.ns import qn
import bs4 # bs4 को ऊपर ले आएं
from kritidocx.utils.style_filter import StyleFilter # फिल्टर को भी ऊपर रखें


logger = logging.getLogger("MyDocX_Engine")

class TableController:
    """
    Main Interface for creating tables.
    Use this class in the Router logic.
    """

    def __init__(self, doc_driver):
        """
        :param doc_driver: Instance of DocxDriver (holding self.doc)
        """
        self.doc = doc_driver.doc

    def _create_and_style_table(self, table_node, container, indent_twips=0):
        """
        [PART 1] Initialization Phase.
        Calculates Geometry, Creates Word Object, and Applies Global Styles.
        
        [UPDATED]: Now supports Raw XML containers (Floating Textboxes).
        """
        # 1. Calculate Geometry (The Matrix)
        matrix, max_cols = MatrixEngine.normalize_structure(table_node)
        
        # Validation: 0-column table Word को क्रैश कर सकती है
        if max_cols < 1:
            logger.warning("Table skipped: No columns detected in structure.")
            return None, None, None

        # 2. Determine Target Container
        target = container if container is not None else self.doc
        word_table = None
        
        # -------------------------------------------------------------
        # 🛠️ [RAW XML FIX START]: Table Creation Strategy
        # -------------------------------------------------------------
        # चेक करें कि क्या कंटेनर 'Standard' (Doc/Cell) है या 'Raw XML' (Textbox)
        # Raw XML एलिमेंट में 'add_table' helper method नहीं होता।
        is_raw_xml = not hasattr(target, 'add_table')

        try:
            # --- [MODIFY "Normal Case" Block inside _create_and_style_table] ---

            if not is_raw_xml:
                # --- Normal Case ---
                try:
                    # 'autofit' property को XML लेवल पर संभालने के लिए हम यहाँ 'width' नहीं देंगे।
                    # Python-docx डिफॉल्ट autofit (rows=0, cols=X) बनाता है जो बेहतर है।
                    word_table = target.add_table(rows=0, cols=max_cols)
                    
                    # Force default style to None to strip ugly Word defaults
                    word_table.style = 'Normal Table' # Or None logic
                except TypeError:
                    # Fallback remains same
                    from docx.shared import Inches
                    word_table = target.add_table(rows=0, cols=max_cols, width=Inches(6.0))
                    
            else:
                # --- Special Case: Floating Box (Raw XML) ---
                # Strategy: Create on Main Doc (temp), Move Node, Reuse Object wrapper
                
                # 1. Create temporary table on body (ताकि एक Valid Python Wrapper मिले)
                temp_table = self.doc.add_table(rows=0, cols=max_cols)
                
                # 2. Move XML Node: (Cut from body, Append to Textbox)
                tbl_element = temp_table._element
                
                # Detach from body (Main Document से हटाना)
                if tbl_element.getparent() is not None:
                    tbl_element.getparent().remove(tbl_element)
                
                # Attach to Textbox content (Floating container में जोड़ना)
                target.append(tbl_element)
                
                # 3. Object Wrapper reuse
                # Python object वही रहेगा, बस उसका इंटरनल XML नये लोकेशन पर चला गया है।
                # यह सुरक्षित है क्योंकि python-docx live XML reference पर काम करता है।
                word_table = temp_table

        except Exception as e:
            # Fallback - Safety Net
            logger.warning(f"Table Create Failed: {e}")
            return None, None, None
        # -------------------------------------------------------------
        # [RAW XML FIX END]
        # -------------------------------------------------------------
                
        # 4. Apply Global Properties (Width, Borders, Alignment, Indent)
        # यह टेबल की बाहरी बनावट तय करता है
        style_str = table_node.get('style', '')
        TablePropsManager.apply_global_styles(word_table, style_str, table_node)

        # 5. Extract Raw HTML Rows (For row-specific styling later)
        # हमें मूल HTML rows भी चाहिए ताकि हम CSS (जैसे height) पढ़ सकें
        original_html_rows = MatrixEngine._extract_linear_rows(table_node)

        # --- Apply List Nesting Indent ---
        # यदि यह तालिका किसी लिस्ट के अंदर है (indent_twips > 0), 
        # तो उसे सही स्थान पर खिसकाएँ।
        if indent_twips > 0:
            from kritidocx.xml_factory.xml_builder import XmlBuilder
            # यह वर्ड के XML में <w:tblInd> सेट करेगा
            XmlBuilder.set_table_indent(word_table, indent_twips)

        return word_table, matrix, original_html_rows
    
    # -------------------------------------------------------------
    # 🛠️ NEW HELPER: RESOLVE BORDER CONFLICTS (MATRIX AWARE)
    # -------------------------------------------------------------
    def _calculate_effective_borders(self, r_idx, c_idx, total_rows, total_cols, cell_css, row_css, table_css):
        """
        Calculates the definitive borders for a specific cell by resolving conflicts
        between Cell, Row, and Table styles.
        
        [UPDATED]: Includes Source-Based Collision Resolution (Collapse Logic).
        """
        final_borders = {}
        
        # डीबग फ्लैग (Config से या सीधे)
        debug = getattr(AppConfig, 'DEBUG_TABLES', False)
        if debug:
             logger.debug(f"\n   [BORDER-LOG] 🎯 Processing Matrix: Row:{r_idx} Col:{c_idx}")

        # --- 1. Helper to Parse ---
        def get_parsed(css_dict, side):
            raw_val = css_dict.get(f'border-{side}') or css_dict.get('border')
            if not raw_val: return None
            
            # यदि पहले से ही Dict है
            if isinstance(raw_val, dict): return raw_val 
            
            # अन्यथा Parse करें
            from kritidocx.basics.border_parser import BorderParser
            return BorderParser.parse(raw_val)

        # --- 2. Priority Resolver with SOURCE TRACKING ---
        # हम केवल बॉर्डर ही नहीं, उसका 'Source' भी वापस करेंगे (Cell vs Row vs Table vs Fallback)
        def resolve_side(side):
            # A. Cell level (Highest Priority)
            b = get_parsed(cell_css, side)
            if b: 
                return b, "CELL"

            # B. Row level
            b = get_parsed(row_css, side)
            if b: 
                return b, "ROW"

            # C. Table Global
            # Note: ग्लोबल बॉर्डर्स अक्सर अंदरूनी लाइनों को कवर नहीं करते, पर यदि defined हैं तो लें
            b = get_parsed(table_css, side)
            if b:
                return b, "TABLE"

            # D. Fallback (Logic Changed)
            # OLD: Return default grid logic
            # NEW: Return None. 
            # अगर कोई स्टाइल नहीं है, तो हम कुछ रिटर्न नहीं करेंगे।
            # इससे Cell Level पर <w:tcBorders> खाली रहेगा और Word 
            # खुद-ब-खुद 'Table Level Grid' (Step 1 वाला) दिखाएगा।
            
            # This PREVENTS the "Single Black vs Double Red" collision conflict.
            return None, "FALLBACK"

        # --- 3. Execute Resolution Loop ---
        perimeter_sides = ['top', 'bottom', 'left', 'right']
        diag_sides = ['tl2br', 'tr2bl']
        
        # यह जानने के लिए कि किस साइड का बॉर्डर कहां से आया
        sources = {}

        for side in perimeter_sides:
            props, source = resolve_side(side)
            
            # [CRITICAL FIX START]: Table या Fallback borders को Cell पर apply न करें.
            # -----------------------------------------------------------------------
            # अगर यह बॉर्डर टेबल की ग्लोबल स्टाइल से आ रहा है या हमारा डिफॉल्ट फॉलबैक है,
            # तो हम इसे final_borders में शामिल नहीं करेंगे।
            # इसका मतलब: सेल XML में <w:tcBorders> खाली रहेगा (उस साइड के लिए)।
            # फायदा: Word नीचे मौजूद Table Grid (w:tblBorders) को दिखाएगा, 
            # जो किसी भी कस्टम सेल बॉर्डर (Double Red) के साथ टकराएगा नहीं, बल्कि उसके पीछे दब जाएगा।
            # -----------------------------------------------------------------------
            
            if source in ["TABLE", "FALLBACK"]:
                sources[side] = source  # Source ट्रैक करें ताकि Grid Logic काम करे (Merge Check के लिए)
                continue                # लेकिन XML प्रॉपर्टीज में मत जोड़ें!
            
            # केवल तभी जोड़ें जब स्रोत 'CELL' या 'ROW' हो
            if props: 
                final_borders[side] = props
                sources[side] = source
            
            if debug:
                logger.debug(f"      ✅ WINNER for {side}: {source} Style")

        
        # Diagonal Logic (Explicit Only)
        for side in diag_sides:
            raw_diag = cell_css.get(f'border-{side}')
            if raw_diag:
                from kritidocx.basics.border_parser import BorderParser
                final_borders[side] = BorderParser.parse(raw_diag)
            else:
                final_borders[side] = {'val': 'nil', 'sz': 0, 'space': 0, 'color': 'auto'}

        # =========================================================
        # 🚀 [COLLISION RESOLUTION - CLEAN GRID STRATEGY]
        # =========================================================
        # समस्या: Word में, यदि पड़ोसी सेल का Right Border है और मेरा Left Border है, 
        # और दोनों एक जैसे हैं, तो कभी-कभी लाइन मोटी दिखती है।
        # समाधान: "Standard Grid" के मामले में 'Left' और 'Top' को दबा दें (Suppress)।
        # (केवल तभी जब हम टेबल के अंदर हों, पहले कॉलम/रो में नहीं)

        # 1. Horizontal Conflict (Suppress LEFT)
        if c_idx > 0:
            # अगर मेरा LEFT बॉर्डर सिर्फ एक "FALLBACK" (साधारण) ग्रिड है, 
            # तो मैं इसे रेंडर नहीं करूंगा। मैं पड़ोसी (Left Cell) के RIGHT बॉर्डर पर भरोसा करूंगा।
            if sources['left'] == 'FALLBACK':
                 final_borders['left'] = {'val': 'nil', 'sz': 0, 'space': 0, 'color': 'auto'}
        
        # 2. Vertical Conflict (Suppress TOP)
        if r_idx > 0:
            # अगर मेरा TOP बॉर्डर सिर्फ एक "FALLBACK" है,
            # तो मैं इसे रेंडर नहीं करूंगा। मैं ऊपर वाले (Upper Cell) के BOTTOM बॉर्डर पर भरोसा करूंगा।
            if sources['top'] == 'FALLBACK':
                 final_borders['top'] = {'val': 'nil', 'sz': 0, 'space': 0, 'color': 'auto'}

        # [ADVANCED]: अगर User ने Explicit 'Cell' बॉर्डर दिया है, तो वह जीतेगा।
        # ऊपर का लॉजिक केवल 'FALLBACK' को हटाता है। अगर 'sources["left"] == "CELL"' है,
        # तो वह 'nil' नहीं होगा और अपनी कस्टम लाइन ड्रॉ करेगा (जैसे Double Red Border)।

        return final_borders

    
    
    def _process_row(self, word_table, row_data, r_idx, original_html_rows, parent_context, router_callback, global_table_styles=None):
        """
        [PART 2] Row Processor.
        Creates the row, applies row-level styles, iterates cells, and executes XML cleanup.
        """
        # 1. Create New Row
        word_row = word_table.add_row()
        tr_element = word_row._element # Raw XML access for cleanup later
        
        # 2. Retrieve HTML Row & Parse Styles
        # (ताकि हम Row Color/Height को Cells में भेज सकें)
        html_tr = original_html_rows[r_idx] if r_idx < len(original_html_rows) else None
        
        # Row Properties (Height, Header Repeat, Page Break)
        RowManager.apply_row_props(word_row, html_tr)
        
        # Extract CSS for inheritance (e.g. background-color red on TR)
        row_styles_dict = {}
        if html_tr:
            from kritidocx.basics.css_parser import CssParser
            row_styles_dict = CssParser.parse(html_tr.get('style', ''))

        # [DEBUG INSERT] - Trace Row Styles
        if getattr(AppConfig, 'DEBUG_TABLES', False):
            # केवल meaningful rows प्रिंट करें
            logger.debug(f"   🔍 [ROW CHECK] Idx:{r_idx} | Parsed Styles: {row_styles_dict}")

        # 3. Track Cells to Delete (The Cleanup List)
        # हम उन सेल्स के इंडेक्स जमा करेंगे जिन्हें XML से हटाना है
        xml_indices_to_delete = []

        # 4. Iterate Through Cells (Matrix Data)
        # Note: word_row.cells में शुरू में उतने ही सेल्स होते हैं जितने टेबल में कॉलम्स हैं (max_cols)
        
        # 1. रो के अंदर मौजूद सभी फिजिकल XML सेल्स की एक स्थिर लिस्ट लें
        from docx.table import _Cell
        all_xml_cells = tr_element.findall(qn('w:tc'))

        # --- [NEW]: टेबल की पूरी साइज निकालें (कोऑर्डिनेट चेक के लिए) ---
        total_rows_count = len(original_html_rows)
        total_cols_count = len(row_data) # रो में मौजूद सेल्स की संख्या

        logger.debug(f"\n🧩 PROCESSING ROW {r_idx}")
        logger.debug(f"   XML Cells Available: {len(all_xml_cells)}")


        for c_idx, cell_info in enumerate(row_data):
            if c_idx >= len(all_xml_cells): break
            
            word_cell = _Cell(all_xml_cells[c_idx], word_table)

            # --- [CRITICAL UPDATE: GHOST STYLE INHERITANCE] ---
            # अगर यह सेल Ghost है (Rowspan का निचला हिस्सा), तो इसका अपना कोई Tag नहीं होता।
            # हमें इसके 'Master' (असली सेल) का स्टाइल लाना होगा ताकि बॉर्डर्स सतत (Continuous) दिखें।
            
            cell_tag = None
            
            # Case 1: यह असली सेल है
            if cell_info.get('type') == 'real':
                cell_tag = cell_info.get('tag')
            
            # Case 2: यह Ghost सेल है (Master से इनहेरिट करें)
            elif cell_info.get('type') == 'merged_placeholder' and 'master' in cell_info:
                # Master Data से tag निकालें
                master_info = cell_info['master']
                cell_tag = master_info.get('tag')

            # अब स्टाइल पार्स करें
            cell_style_str = cell_tag.get('style', '') if cell_tag else ''
            
            from kritidocx.basics.css_parser import CssParser
            cell_css_dict = CssParser.parse(cell_style_str)
            # --------------------------------------------------------


            # 💡 मुख्य लॉजिक कॉल: यह तय करेगा कि (Black Grid) आएगा या (Custom Wavy Red)
            effective_border_map = self._calculate_effective_borders(
                r_idx=r_idx, 
                c_idx=c_idx,
                total_rows=total_rows_count, 
                total_cols=total_cols_count,
                cell_css=cell_css_dict, 
                row_css=row_styles_dict,
                table_css=global_table_styles or {}
            )
            # --- [सुधार END] ---
            
            # अब यह map 'calculated_borders' के रूप में आगे भेजें
            indices_to_remove = self._process_individual_cell(
                word_cell, 
                cell_info, 
                c_idx, 
                router_callback, 
                row_styles_dict, 
                parent_context,
                global_table_styles=global_table_styles,
                calculated_borders=effective_border_map # <--- यह नया पैरामीटर भेजें
            )
            
            if indices_to_remove:
                # लिस्ट में हम जोड़ रहे हैं, इसलिए 'extend' यूज करें ताकि सभी जमा हो जाएं
                xml_indices_to_delete.extend(indices_to_remove) 
                logger.debug(f"   ⚡ Cell {c_idx} requested deleting indices: {indices_to_remove}")
  
            
            
        # =========================================================
        # 5. FINAL CLEANUP PHASE (The Nuclear Fix)
        # =========================================================
        if xml_indices_to_delete:
            # --- [NEW LOG] ---
            logger.debug(f"\n   🛠️ CLEANUP START for Row {r_idx}:")
            logger.debug(f"      Initial Request List: {xml_indices_to_delete}")

            # 1. यूनिक और रिवर्स सॉर्टेड इंडेक्स
            sorted_indices = sorted(list(set(xml_indices_to_delete)), reverse=True)
            
            # --- [NEW LOG] ---
            logger.debug(f"      Unique Sorted (Reverse): {sorted_indices}")

            # 2. दोबारा ताज़ा XML लिस्ट लें
            current_tcs = tr_element.findall(qn('w:tc'))
            logger.debug(f"      Physical XML Cells actually present: {len(current_tcs)}")

            for del_idx in sorted_indices:
                if del_idx < len(current_tcs):
                    # 3. फिजिकल नोड को डिलीट करें
                    node_to_remove = current_tcs[del_idx]
                    tr_element.remove(node_to_remove)
                    if getattr(AppConfig, 'DEBUG_TABLES', False):
                        logger.debug(f"      ✅ Deleted XML Cell at index: {del_idx}")

                else:
                    if getattr(AppConfig, 'DEBUG_TABLES', False):
                        logger.debug(f"      ⚠️ SKIP: index {del_idx} out of range!")

        else:
            # --- [NEW LOG] ---
            logger.debug(f"   ✨ Row {r_idx}: No cleanup required.")   
           
           
                       
    def _process_individual_cell(self, word_cell, cell_info, c_idx, router_callback, row_styles, parent_context, global_table_styles=None,calculated_borders=None):
        """
        [PART 3] Cell Logic Unit.
        Decides the fate of a cell based on Matrix Data.
        
        Returns:
            list: Indices of XML cells to delete (e.g. [1, 2] if colspan=3).
        """
        indices_to_remove = []


        # =========================================================
        # 🛡️ MERGE BORDER SURGERY (HTML Like Box Fix)
        # =========================================================
        # समस्या: Rowspan के बीच में Word डिफ़ॉल्ट रूप से बॉर्डर लगा देता है।
        # समाधान: Restart का Bottom और Continue का Top हटा दें।
        if calculated_borders:
            merge_state = cell_info.get('v_merge')
            
            # Case 1: ऊपरी हिस्सा (Start) -> नीचे का बॉर्डर हटाएं ताकि वह नीचे वाले सेल से मिल जाए
            if merge_state == 'restart':
                calculated_borders['bottom'] = {'val': 'nil', 'sz': 0, 'space': 0, 'color': 'auto'}
            
            # Case 2: निचला हिस्सा (Continuation) -> ऊपर का बॉर्डर हटाएं
            elif merge_state == 'continue':
                calculated_borders['top'] = {'val': 'nil', 'sz': 0, 'space': 0, 'color': 'auto'}
        # =========================================================


        # --- CASE 1: EMPTY PAD CELL ---
        # (Matrix Engine ने इसे रेक्टैंगल पूरा करने के लिए जोड़ा था)
        if cell_info is None or cell_info['type'] == 'pad':
            return []

        # --- CASE 2: REAL DATA CELL ---
        elif cell_info['type'] == 'real':
            # A. Vertical Merge Start (Rowspan)
            if cell_info.get('v_merge') == 'restart':
                from kritidocx.xml_factory.xml_builder import XmlBuilder
                XmlBuilder.set_cell_v_merge(word_cell, "restart")

            # B. Horizontal Merge (Colspan)
            g_span = cell_info.get('grid_span')
            if g_span and g_span > 1:
                from kritidocx.xml_factory.xml_builder import XmlBuilder
                XmlBuilder.set_cell_grid_span(word_cell, g_span)
                
                # [CRITICAL LOGIC]: अगर मैं 2 कॉलम घेर रहा हूँ, तो मेरे बगल वाला सेल हटना चाहिए।
                # हम उसे 'indices_to_remove' में नोट कर लेते हैं।
                for k in range(1, g_span):
                    indices_to_remove.append(c_idx + k)

            # C. Content & Styling (Delegate to CellManager)
            # यहाँ हम असली कंटेंट भरते हैं
            CellManager.process_cell(
                word_cell, 
                cell_info, 
                router_callback=router_callback,
                row_styles=row_styles,
                table_defaults=parent_context,
                table_global_styles=global_table_styles,
                calculated_borders=calculated_borders
            )

        # --- CASE 3: MERGED GHOST CELL (The Invisible Part) ---
        elif cell_info['type'] == 'merged_placeholder':
            # Sub-Case A: Horizontal Ghost (Colspan का हिस्सा)
            # इसे तुरंत डिलीट करना है, चाहे यह rowspan का हिस्सा हो या नहीं
            if cell_info.get('is_h_merged'):
                indices_to_remove.append(c_idx)
                return indices_to_remove

            # Sub-Case B: Vertical Ghost (Rowspan का हिस्सा)
            if cell_info.get('v_merge') == 'continue':
                from kritidocx.xml_factory.xml_builder import XmlBuilder
                XmlBuilder.set_cell_v_merge(word_cell, "continue")
                
                # [FIX]: मास्टर की चौड़ाई चेक करें
                if 'master' in cell_info:
                    m_span = cell_info['master'].get('grid_span')
                    if m_span and m_span > 1:
                        # मास्टर की चौड़ाई (gridSpan) यहाँ भी लागू करें
                        XmlBuilder.set_cell_grid_span(word_cell, m_span)
                        # इस रो में इसके बगल वाले एक्स्ट्रा सेल्स को डिलीट करें
                        for k in range(1, m_span):
                            indices_to_remove.append(c_idx + k)

                # [COLOR FIX] Ghost cells को मास्टर का स्टाइल पास करें
                # इससे 'ACTIVE' (Green Box) पूरा नीचे तक हरा दिखेगा
                ghost_styles = row_styles.copy() if row_styles else {}
                
                if 'master' in cell_info:
                    master_node = cell_info['master'].get('tag')
                    if master_node:
                        # मास्टर सेल की इनलाइन स्टाइल (जैसे background-color) पार्स करें
                        from kritidocx.basics.css_parser import CssParser
                        m_styles = CssParser.parse(master_node.get('style', ''))
                        ghost_styles.update(m_styles)

                # CellManager को कॉल करें (Style only, no content)
                CellManager.process_cell(
                    word_cell, 
                    cell_info, 
                    router_callback=None, 
                    row_styles=ghost_styles ,# अपडेटेड स्टाइल भेजें
                    calculated_borders=calculated_borders
                )


        return indices_to_remove
    
    
    def process_table(self, table_node, container=None, parent_context=None, router_callback=None, indent_override=0):
        """
        [PART 4] Main Orchestrator (The Manager).
        Converts HTML <table> into Word Table using the Refactored Pipeline.
        
        Flow:
        1. Initialization (Geometry & Styles) -> _create_and_style_table
        2. Row Iteration (Processing & Cleanup) -> _process_row
        3. Cell Logic (Merging & Content) -> _process_individual_cell (called by row)
        """
        if not table_node: return None

        # -------------------------------------------------------------
        # STEP 1: INITIALIZATION PHASE
        # -------------------------------------------------------------
        # यह Matrix बनाता है, Word Table ऑब्जेक्ट बनाता है और Global Styles लगाता है
        word_table, matrix, original_html_rows = self._create_and_style_table(table_node, container, indent_twips=indent_override)

        # Validation: अगर टेबल नहीं बनी (जैसे 0 कॉलम्स), तो रुक जाएं
        if not word_table:
            return None

        # टेबल का मूल स्टाइल स्ट्रिंग निकालें
        style_str = table_node.get('style', '')
        
        # इसे डिक्शनरी में बदलें ताकि हम background-color पढ़ सकें
        from kritidocx.basics.css_parser import CssParser
        table_style_dict = CssParser.parse(style_str)

        # -------------------------------------------------------------
        # STEP 2: ROW PROCESSING LOOP
        # -------------------------------------------------------------
        # अब हम मैट्रिक्स की हर पंक्ति (Row) पर लूप चलाएंगे
        for r_idx, row_data in enumerate(matrix):
            
            # सारी जटिलता (Complexity) को _process_row को सौंप दें
            # यह पंक्ति बनाएगा, सेल्स भरेगा, और एक्स्ट्रा सेल्स को डिलीट (Cleanup) करेगा
            self._process_row(
                word_table=word_table,
                row_data=row_data,
                r_idx=r_idx,
                original_html_rows=original_html_rows,
                parent_context=parent_context,
                router_callback=router_callback,
                global_table_styles=table_style_dict
            )

        # -------------------------------------------------------------
        # STEP 3: FINAL POST-PROCESSING
        # -------------------------------------------------------------
        # (Optional) नेस्टेड टेबल्स के लिए एक्स्ट्रा स्पेसिंग लॉजिक यहाँ आ सकता है
        # लेकिन अभी इसकी जरूरत नहीं है क्योंकि CellManager इसे संभाल रहा है।

        return word_table
    
    def create_box_container(self, styles, container):
        """
        DIV बॉर्डर को 1x1 वर्ड टेबल बॉक्स में बदलता है।
        [MERGED FEATURES]: 
        1. Dynamic Width: पेज साइज (A4, Landscape) के अनुसार खुद को एडजस्ट करता है।
        2. Safe Layout: Header/Footer में क्रैश नहीं होता।
        3. Formatting: बाहरी स्पेस और अंदरूनी लाइन हाइट को संतुलित करता है।
        """
        from docx.shared import Pt, Inches, Twips
        from docx.enum.text import WD_LINE_SPACING
        
        from kritidocx.basics.unit_converter import UnitConverter
        from kritidocx.xml_factory.xml_builder import XmlBuilder
        from kritidocx.xml_factory.table_xml import TableXml
        from .cell_manager import CellManager

        # --- 🛡️ 1. Smart Target & Width Calculation ---
        # दस्तावेज़ के अंतिम सेक्शन से पेज की चौड़ाई और मार्जिन निकालें
        section = self.doc.sections[-1]
        
        # डिफ़ॉल्ट मान (सुरक्षा के लिए, अगर सेक्शन डेटा न मिले)
        page_w = section.page_width or Twips(12240) # ~8.5 inch
        left_m = section.left_margin or Twips(1440) 
        right_m = section.right_margin or Twips(1440)
        
        # असली लिखने योग्य जगह (Printable Area Calculation)
        # 50 Twips का बफर रखें ताकि बॉर्डर मार्जिन से बाहर न गिरे
        available_width_val = page_w - left_m - right_m - Twips(50) 
        
        # लक्ष्य तय करें (Body vs Container)
        target = container if container is not None else self.doc
        if not hasattr(target, 'add_table'):
            if hasattr(target, '_parent') and hasattr(target._parent, 'add_table'):
                target = target._parent
            else:
                target = self.doc

        # --- 🛡️ 2. Safe & Responsive Table Creation ---
        try:
            # प्रयास 1: साधारण तरीका (Standard Doc/Cell)
            box_table = target.add_table(rows=1, cols=1)
            
        except TypeError as e:
            # प्रयास 2: अगर Header/Footer 'width' मांगता है (Fallback)
            if 'width' in str(e) or 'required positional argument' in str(e):
                # यहाँ हम फिक्स्ड '6.0 inch' के बजाय 'available_width_val' का उपयोग करेंगे
                box_table = target.add_table(rows=1, cols=1, width=available_width_val)
            else:
                logger.error(f"   ⚠️ Table Add TypeError inside Box: {e}")
                box_table = self.doc.add_table(rows=1, cols=1)
                
        except Exception as e:
            # प्रयास 3: पूर्ण विफलता -> Body Fallback
            logger.error(f"   ⚠️ Generic Table Error inside Box: {e} -> Fallback to Doc Body")
            box_table = self.doc.add_table(rows=1, cols=1)

        cell = box_table.cell(0, 0)

        # --- 3. Base Formatting Application ---
        CellManager.process_cell(
            word_cell=cell, 
            cell_info={'type': 'real', 'tag': None},
            router_callback=None, 
            row_styles=styles
        )

        # --- 4. Advanced Positioning & Indentation ---
        
        # A. Indentation Logic (Left Margin / Indent)
        m_left = styles.get('margin-left') or styles.get('margin_left')
        
        # यह ट्रैक करेगा कि क्या टेबल को शिफ्ट किया गया है?
        active_indent_val = 0 
        
        if m_left and m_left != 'auto':
            active_indent_val = UnitConverter.to_twips(str(m_left))
            if active_indent_val > 0:
                XmlBuilder.set_table_indent(box_table, active_indent_val)
                # यदि इंडेंट दिया गया है, तो Width ऑटोमैटिक (Shrink) मोड में जा सकती है (जब तक explicit width न हो)
                if 'width' not in styles: 
                    XmlBuilder.set_table_width_pct(box_table, 0)

        # B. Alignment Logic (Auto Margin)
        m_right = styles.get('margin-right')
        if m_left == 'auto' and m_right == 'auto':
            XmlBuilder.set_table_alignment(box_table, 'center')

        # --- 5. ULTRA-ACCURATE WIDTH ENGINE (Fixed + Dynamic Merged) ---
        explicit_width = styles.get('width')
        is_actually_shifted = (active_indent_val > 0)

        if explicit_width:
            # Scenario A: यूजर ने चौड़ाई (Width) CSS में दी है
            raw_w = str(explicit_width).strip()
            if '%' in raw_w:
                # Percentage (%)
                pct_val = UnitConverter.to_table_pct(raw_w)
                XmlBuilder.set_table_width_pct(box_table, pct_val)
                # % को ठीक से काम करने के लिए fixed layout ज़रूरी है
                TableXml.set_table_layout(box_table, 'fixed')
            else:
                # Fixed Units (px, in, cm)
                twips_val = UnitConverter.to_twips(raw_w)
                if twips_val > 0:
                    TableXml.set_table_width(box_table, str(twips_val), 'dxa')
        
        elif not is_actually_shifted:
            # Scenario B: Auto Full Width (Smart Expansion)
            # अगर कोई Width नहीं और कोई Indent नहीं, तो इसे पूरे पेज पर फैलाओ।
            
            # 1. XML width property को 100% (5000 pct) सेट करें
            XmlBuilder.set_table_width_pct(box_table, 5000)
            
            # 2. 🔥 NUCLEAR FIX: अंतर्निहित ग्रिड (Table Grid) को Override करें
            # यह सुनिश्चित करता है कि हेडर A4 और Landscape दोनों पर बिना कटे फिट हो।
            grid_col_width = int(available_width_val)
            TableXml.define_table_grid(box_table, [grid_col_width])
            
            # 3. लेआउट लॉक करें
            TableXml.set_table_layout(box_table, 'fixed')

        # --- 6. Internal Content Polish (Cleanups) ---
        
        # पैडिंग चेक: क्या CSS में पैडिंग दी गई है?
        has_css_padding = any(k for k in styles.keys() if 'padding' in k)

        # (i) Smart Cell Margins
        if not has_css_padding:
            # अगर पैडिंग नहीं है, तो सेल मार्जिन को 'tight' करें लेकिन टेक्स्ट बॉर्डर से न चिपके
            XmlBuilder.set_cell_margins(cell, {'top': 0, 'bottom': 0, 'left': 100, 'right': 100})
        else:
            # पैडिंग होने पर CellManager के मान ही रहने दें
            pass

        # (ii) Paragraph Spacing & Fidelity
        if cell.paragraphs:
            p = cell.paragraphs[0]
            
            # अंदरूनी अनचाही खाली जगह हटाएँ
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            if has_css_padding:
                # पैडिंग होने पर लाइन स्पेसिंग 'Single' (सुरक्षित) रखें
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            else:
                # पैडिंग नहीं होने पर (Compact Border), इसे टाइट रखें (1pt Exact)
                p.paragraph_format.line_spacing = Pt(1) 
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        # --- 7. External Spacing (Buffer) ---
        # यह बॉक्स के बाद एक छोटा खाली पैराग्राफ जोड़ता है ताकि दो टेबल आपस में चिपके नहीं।
        if hasattr(target, 'add_paragraph'):
            after_spacer = target.add_paragraph()
            after_spacer.paragraph_format.space_before = Pt(2)
            after_spacer.paragraph_format.space_after = Pt(2)
            after_spacer.paragraph_format.line_spacing = Pt(1)

        return cell
    
    
    
    def create_flex_layout_grid(self, node, container, context, router_callback):
        """
        [ULTIMATE VERSION]: Full-Stack Architecture Compliance.
        Responsibilities:
        1. [Table Layer] Create Grid Structure & Dimensions (Fixed/Autofit logic).
        2. [Style Layer] Apply Backgrounds with proper Parent/Child inheritance logic.
        3. [Router Layer] Delegate content processing back to Core Router.
        4. [Clean Layer] Strip margins/padding to prevent wrapping issues.
        """
        # LOCAL IMPORTS
        from kritidocx.basics.css_parser import CssParser
        from kritidocx.xml_factory.xml_builder import XmlBuilder
        from kritidocx.xml_factory.table_xml import TableXml
        from kritidocx.utils.style_filter import StyleFilter
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import bs4

        # 1. Child Analysis (Get Nodes)
        children = [c for c in node.children if isinstance(c, bs4.element.Tag) and c.name not in ['script', 'style']]
        if not children:
            return False

        col_count = len(children)
        from docx.section import _Header, _Footer # ये पता करने के लिए कि हम कहाँ हैं

        # [NEW LOGIC]: स्मार्टली तय करें कि 'width' देना है या नहीं
        # हम सिर्फ Header या Footer ऑब्जेक्ट्स में 'width' एर्ग्यूमेंट पास करेंगे।
        is_in_hf = isinstance(container, (_Header, _Footer))

        # -------------------------------------------------------------
        # PHASE 1: STRUCTURE (Smart Context-Aware Grid Creation)
        # -------------------------------------------------------------
        
        # [FIX 1]: डायनामिक चौड़ाई (Dynamic Width Calculation)
        # 6.0 इंच हार्डकोडिंग के बजाय पेज के मार्जिन से गणना करें।
        
        target_width_obj = None
        
        if is_in_hf:
            try:
                # पेज की सेटिंग निकालें
                section = self.doc.sections[-1]
                page_w = section.page_width or Twips(11906) # Default A4 Width
                left_m = section.left_margin or Twips(1440) # Default 1 inch
                right_m = section.right_margin or Twips(1440)
                
                # वास्तविक लिखने योग्य चौड़ाई निकालें (thoda buffer hata kar safe side ke liye)
                # हम 100 twips का बफर कम करेंगे ताकि बॉर्डर कटे नहीं
                calc_width = page_w - left_m - right_m - Twips(50)
                target_width_obj = calc_width
            except:
                # अगर गणित फेल हो जाए, तो सुरक्षित 6.2 इंच लें
                target_width_obj = Inches(6.2)

        try:
            if is_in_hf and target_width_obj:
                # [FIX APPLIED]: Pass Dynamic Width
                table_obj = container.add_table(rows=1, cols=col_count, width=target_width_obj)
            else:
                # Body Normal Flow
                table_obj = container.add_table(rows=1, cols=col_count)
            
            # ग्रिड लाइन्स और बॉर्डर हटाना
            from kritidocx.xml_factory.table_xml import TableXml
            TableXml.set_table_borders_to_none(table_obj)
            
            # चौड़ाई 100% (5000 pct units)
            TableXml.set_table_width(table_obj, 5000, 'pct')
            TableXml.set_table_layout_preset(table_obj, 'autofit')
              
        except TypeError as e:
            # एक्स्ट्रा सेफ्टी: अगर ऊपर वाला डिसीजन भी गलत हो जाए तो ये अंतिम बैकअप है
            logger.debug(f"🔄 Typing mismatch detected, retrying fallback table creation...")
            table_obj = container.add_table(rows=1, cols=col_count)
            
        except Exception as e:
            logger.error(f"⚠️ Fatal Table Creation Failure: {e}")
            return False


        # Access Cells
        cells = table_obj.rows[0].cells
        
        # -------------------------------------------------------------
        # PHASE 2: PROCESSING LOOP (Style & Content)
        # -------------------------------------------------------------
        for i, child_node in enumerate(children):
            cell = cells[i]
            
            # --- STYLE PARSING ---
            c_style = CssParser.parse(child_node.get('style', ''))
            
            # [FIX 2]: Respect HTML Padding (CSS Padding को मानें)
            # पुराना कोड सीधे reset कर रहा था: top=0, bottom=0...
            # नया कोड: context (पैरेंट) की पैडिंग चेक करें, यदि मौजूद है तो लगाएं
            
            # CSS पैडिंग पार्स करें
            pad_left = context.get('padding-left') or '40' # default tiny space
            pad_right = context.get('padding-right') or '40'
            pad_top = context.get('padding-top') or '0'
            pad_bottom = context.get('padding-bottom') or '0'
            
            # Convert units via helper (यदि helper उपलब्ध नहीं तो logic सरल रखें)
            from kritidocx.basics.unit_converter import UnitConverter
            
            # UnitConverter का उपयोग करके सुरक्षित Twips निकालें (CSS '20px' -> ~300 Twips)
            # यदि helper call महँगा है, तो direct mapping का उपयोग करें
            pl_twips = UnitConverter.to_twips(str(pad_left)) if pad_left else 40
            pr_twips = UnitConverter.to_twips(str(pad_right)) if pad_right else 40
            pt_twips = UnitConverter.to_twips(str(pad_top)) if pad_top else 0
            pb_twips = UnitConverter.to_twips(str(pad_bottom)) if pad_bottom else 0

            # अब अप्लाई करें
            XmlBuilder.set_grid_cell_margins(
                cell, 
                top=pt_twips, 
                bottom=pb_twips, 
                left=pl_twips, 
                right=pr_twips
            )
            
            # --- [SMART WIDTH LOGIC (Flex Basis)] ---
            # अगर यूज़र ने CSS में चौड़ाई (Width) दी है, तो उसे सेल पर लागू करें
            child_width = c_style.get('width')
            if child_width:
                from kritidocx.basics.unit_converter import UnitConverter
                if '%' in str(child_width):
                    # '40%' को Word के Pct फॉर्मेट में बदलें
                    pct_val = UnitConverter.to_table_pct(str(child_width))
                    XmlBuilder.set_cell_width(cell, str(pct_val), 'pct')
                else:
                    # '200px' जैसे फिक्स्ड यूनिट्स के लिए
                    twips_val = UnitConverter.to_twips(str(child_width))
                    XmlBuilder.set_cell_width(cell, str(twips_val), 'dxa')
            
            # (नोट: अगर यूज़र ने चौड़ाई नहीं दी है, तो कुछ सेट न करें। 
            # वर्ड का 'Autofit' इंजन टेक्स्ट के आधार पर इसे खुद सिकुड़/फैला लेगा!)         
            
            # Vertical Align (Top preferred for headers)
            valign = c_style.get('vertical-align', 'top')
            XmlBuilder.set_cell_valign(cell, valign)

            # --- [BACKGROUND INHERITANCE FIX] ---
            # 1. चाइल्ड (Cell) का कलर चेक करें
            child_bg = c_style.get('background-color') or c_style.get('background')
            # 2. पैरेंट (Row/Container) का कलर चेक करें
            parent_bg = context.get('background-color') or context.get('background')
            
            # लॉजिक: बच्चे का कलर ले लो, नहीं तो बाप का कलर ले लो (Hierarchy)
            final_bg = child_bg if child_bg else parent_bg
            
            if final_bg and str(final_bg).lower() not in ['transparent', 'none']:
                from kritidocx.basics.color_manager import ColorManager
                hex_val = ColorManager.get_hex(final_bg)
                if hex_val: 
                    XmlBuilder.set_cell_shading(cell, hex_val)

            # --- [BORDER HANDLING] ---
            # सेल लेवल के बॉर्डर को प्राथमिकता दें
            from kritidocx.objects.table.cell_manager import CellManager
            CellManager._apply_borders(cell, c_style)

            # -------------------------------------------------------------
            # PHASE 3: CONTENT DELEGATION (Router Call)
            # -------------------------------------------------------------
            if router_callback:
                # Context Sanitation (Stop leaking block props to children)
                child_context = StyleFilter.get_clean_child_context(context)
                
                # 1. Determine Horizontal Alignment
                target_align = c_style.get('text-align')
                justify_val = context.get('justify-content', '')
                # Auto-alignment logic (Space Between)
                # अगर 2 कॉलम हैं तो [Left, Right]. अगर 3 हैं तो [Left, Center, Right]
                if not target_align and 'space-between' in justify_val:
                    if i == 0: 
                        target_align = 'left'
                    elif i == col_count - 1: 
                        target_align = 'right'
                    else: 
                        target_align = 'center'
                
                # Context Update
                if target_align:
                    child_context['text-align'] = target_align
                    child_context['align'] = target_align

                # 2. Preparation (Emptying default paragraph)
                if len(cell.paragraphs) > 0:
                    p = cell.paragraphs[0]
                    p.text = "" # साफ़ करें
                else:
                    p = cell.add_paragraph()

                # 3. Router Recursion
                # हम सेल के अंदर मौजूद सभी 'Grandchildren' (H1, Span, Text) को Router को देते हैं
                content_found = False
                for inner_child in child_node.children:
                    if isinstance(inner_child, bs4.element.NavigableString) and not str(inner_child).strip():
                        continue # Skip pure whitespace
                        
                    # ⚠️ Critical: Pass CELL as container, NOT paragraph. 
                    # Router will decide whether to reuse the existing para or make new block elements.
                    router_callback(inner_child, cell, child_context)
                    content_found = True

                # 4. Final Alignment Polish
                # राउटर के काम पूरा करने के बाद, हम सुनिश्चित करते हैं कि सभी पैराग्राफ
                # सही दिशा में अलाइन हों (Right/Center/Left) और मार्जिन टाइट हों।
                if content_found and target_align:
                    wd_align = None
                    if 'right' in target_align: wd_align = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    elif 'center' in target_align: wd_align = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif 'justify' in target_align: wd_align = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    else: wd_align = WD_PARAGRAPH_ALIGNMENT.LEFT
                    
                    if wd_align is not None:
                        for p in cell.paragraphs:
                            # अलाइनमेंट सेट करें
                            p.alignment = wd_align
                            # Header में एक्स्ट्रा स्पेस हटाने के लिए
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)

        # Process complete
        return True