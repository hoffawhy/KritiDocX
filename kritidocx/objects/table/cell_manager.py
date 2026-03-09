"""
from kritidocx.utils.logger import logger
CELL MANAGER MODULE (The Grid Decorator)
----------------------------------------
Responsibility:
Handles formatting, merging logic, and content filling for individual Table Cells.

Key Features:
1. Inheritance Model: Merges Inline Styles with Row Styles.
2. Geometry Logic: Converts Rowspan/Colspan to XML 'vMerge'/'gridSpan'.
3. Styling: Borders, Background Shading, Vertical Alignment, Rotation.
4. Recursion Bridge: Calls the Router to process content INSIDE the cell.

Word Specifics:
- An HTML 'rowspan' becomes 'restart' in first cell, 'continue' in subsequent hidden cells.
"""

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn


from kritidocx.config.settings import AppConfig
from kritidocx.objects.text.paragraph_manager import ParagraphManager 
from kritidocx.basics.css_parser import CssParser
from kritidocx.basics.unit_converter import UnitConverter
from kritidocx.basics.color_manager import ColorManager
from kritidocx.basics.border_parser import BorderParser
from kritidocx.basics.font_handler import FontHandler
from kritidocx.utils import logger
from kritidocx.xml_factory.xml_builder import XmlBuilder
from kritidocx.config.theme import ThemeConfig

class CellManager:
    """
    Controller for operations on a Single Table Cell.
    """

    @classmethod
    def process_cell(cls, word_cell, cell_info, router_callback, row_styles=None, table_defaults=None, table_global_styles=None,calculated_borders=None):
        """
        Master method to setup a single cell in the Word Table Grid.
        
        Args:
            word_cell: python-docx _Cell object.
            cell_info (dict): Metadata from MatrixEngine (tag, rowspan, type).
            router_callback (func): function(node, container, style) to handle recursing children.
            row_styles (dict): Styles inherited from the parent <TR>.
            table_defaults (dict): Defaults for the whole table.
        """
        if not cell_info: return # Empty/Spacer cell

        # जब हम नया सेल बनाते हैं, तो उसमें डिफ़ॉल्ट <w:p/> पहले से मौजूद होता है।
        # उसे हटा दें, ताकि हमारा कंटेट टॉप से शुरू हो, न कि 1 लाइन नीचे से।
        # (नोट: Word को कम से कम 1 पैराग्राफ चाहिए होता है, लेकिन हम नीचे अपना बनाएंगे, इसलिए इसे हटाना सुरक्षित है)
        
            
        # =========================================================
        # 1. IDENTIFY STATE (Merged vs Real)
        # =========================================================
        is_ghost = (cell_info['type'] == 'merged_placeholder')
        
        # Merge Flags (Word XML)
        # Vertical Merge logic is tricky. 
        # 'restart' = Starts span. 'continue' = Hidden part of span.
        v_merge_val = cell_info.get('v_merge') # 'restart', 'continue', or None
        
        if v_merge_val:
            XmlBuilder.set_cell_v_merge(word_cell, v_merge_val)

        # Handle 'continue' ghosts (Merged placeholders)
        # We process them lightly just to ensure borders continuity if needed
        # but usually we skip content filling.
        if is_ghost:
            if 'master' in cell_info and not cell_info.get('is_h_merged'):
                # --- [FIX START] ---
                # Prioritize explicit calculated borders from TableController logic
                if calculated_borders:
                    XmlBuilder.set_cell_borders(word_cell, calculated_borders)
                    
                    # shading logic remains manual for ghosts if needed
                    # (Usually extract from master tag logic inside helper, 
                    # but safer to decouple border logic)
                    cls._apply_ghost_styling_shading_only(word_cell, cell_info['master'], row_styles)
                else:
                    # Fallback to old behavior
                    cls._apply_ghost_styling(word_cell, cell_info['master'], row_styles)
                # --- [FIX END] ---
            return

        # =========================================================
        # 2. REAL CELL PROCESSING
        # =========================================================
        # केवल 'Real' सेल के पैराग्राफ डिलीट करें ताकि ऊपर स्पेस न आये
        if cell_info['type'] == 'real':
            for p in word_cell.paragraphs:
                p._element.getparent().remove(p._element)

     
        html_node = cell_info['tag']
        
        # अगर tag 'None' है (Box Model के केस में), तो डिफ़ॉल्ट 'td' मानें
        tag_name = html_node.name.lower() if html_node else 'td'
        
        # Parse CSS (Safety Check for node)
        style_str = html_node.get('style', '') if html_node else ''
        inline_styles = CssParser.parse(style_str)

        # यह चेक करता है कि क्या यूजर ने 'white-space: nowrap' दिया है।
        # अगर हाँ, तो सेल की चौड़ाई कंटेंट के हिसाब से फैलेगी, टूटेगी नहीं।
        white_space_val = inline_styles.get('white-space', '').lower()
        if white_space_val in ['nowrap', 'pre']:
            # XmlBuilder.set_cell_no_wrap का इस्तेमाल करें (जिसे आपने अभी बनाया है)
            XmlBuilder.set_cell_no_wrap(word_cell, True)

        
        # A. Horizontal Merging (gridSpan)
        colspan = cell_info.get('colspan', 1)
        if colspan > 1:
            XmlBuilder.set_cell_grid_span(word_cell, colspan)

        # B. Formatting (The heavy lifting)
        # Context Style created here is passed down to children text
        context_style = cls._apply_visuals(
            word_cell, 
            html_node, 
            inline_styles, 
            row_styles,
            is_header=(tag_name == 'th'),
            table_styles=table_global_styles,
            explicit_border_map=calculated_borders
        )

        # =========================================================
        # 3. FILL CONTENT (Smart Inline Grouping)
        # =========================================================
        if router_callback and html_node:
            
            # वह HTML टैग्स जो 'Block Level' माने जाते हैं और नया पैराग्राफ माँगते हैं
            BLOCK_TAGS = ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 
                        'ul', 'ol', 'table', 'blockquote', 'pre', 'hr', 'form']

            # 'active_paragraph' वह पैराग्राफ है जिसमें हम इनलाइन टेक्स्ट भरेंगे
            active_paragraph = None
            content_found = False
            
            import bs4

            for child in html_node.children:
                
                # --- A. Whitespace/Empty Text को इग्नोर करें ---
                if isinstance(child, bs4.element.NavigableString):
                    if not str(child).strip(' \n\r\t\u200b'): 
                        continue

                
                # --- B. टैग का प्रकार पता करें ---
                tag_name = child.name.lower() if isinstance(child, bs4.element.Tag) else None
                is_block_element = tag_name in BLOCK_TAGS

                # --- C. निर्णय लें: एक साथ रखें या अलग करें? ---
                
                if is_block_element:
                    # 1. ब्लॉक एलीमेंट (जैसे <p>): इसके लिए "Active Paragraph" तोड़ दें
                    # राउटर को 'word_cell' दें ताकि वह खुद नया पैराग्राफ बना सके
                    router_callback(child, word_cell, context_style)
                    active_paragraph = None # रिसेट करें
                    
                else:
                    # 2. इनलाइन एलीमेंट (Text, Span, Bold, BR): एक पैराग्राफ में जोड़ें
                    
                    # अगर कोई चालू पैराग्राफ नहीं है, तो नया बनाएं और उस पर स्टाइल (0pt margin) लगाएं
                    if active_paragraph is None:
                        active_paragraph = word_cell.add_paragraph()
                        
                        #  Apply Tight Layout immediately to container paragraph
                        ParagraphManager.apply_formatting(active_paragraph, context_style)

                    # राउटर को 'active_paragraph' दें ताकि वह content को इसी में append करे (Soft Break logic)
                    router_callback(child, active_paragraph, context_style)

        # =========================================================
        # 4.  FINAL SAFETY: EMPTY CELL POLICY
        # =========================================================
        # हम जांचेंगे कि क्या कंटेंट भरने के बाद भी सेल 'प्रभावी रूप से' खाली है?
        # (जैसे: <td></td> या <td>   </td>)
        
        is_effectively_empty = True

        # पैराग्राफ्स में लूप लगाकर असली कंटेंट ढूँढें
        for p in word_cell.paragraphs:
            # चेक करें: टेक्स्ट है? या कोई रन (Format/Symbol)? या कोई इमेज/ड्राइंग?
            if p.text.strip() or p.runs or p._element.find(qn('w:drawing')) is not None:
                is_effectively_empty = False
                break
        
        if is_effectively_empty:
            # CASE: Cell is completely empty -> Prevent Collapse
            
            # एक पैराग्राफ पकड़ें (या बनाएं)
            if len(word_cell.paragraphs) == 0:
                target_p = word_cell.add_paragraph()
            else:
                target_p = word_cell.paragraphs[0]
            
            # 'Zero Width Space' (\u200b) डालें
            # यह अदृश्य है लेकिन Word को बताता है कि लाइन की हाइट कायम रखें।
            target_p.add_run('\u200b') 
            
            # मार्जिन को 0 करें ताकि सेल फूला हुआ न दिखे
            p_fmt = target_p.paragraph_format
            p_fmt.space_before = Pt(0)
            p_fmt.space_after = Pt(0)
            
            # Context styles (Background color etc.) लागू करें
            ParagraphManager.apply_formatting(target_p, context_style)

        else:
            # CASE: Cell has content -> Remove Trailing Spacer
            # अगर कंटेंट है, तो अक्सर आखिरी पैराग्राफ एक एक्स्ट्रा खाली लाइन होती है। उसे हटा दें।
            cls._cleanup_trailing_space(word_cell)

           
    # -------------------------------------------------------------
    # 🖌️ VISUAL STYLING ENGINE
    # -------------------------------------------------------------

    @classmethod
    def _apply_visuals(cls, cell, node, styles, row_styles, is_header, table_styles=None,explicit_border_map=None):
        """
        Applies Borders, Colors, Alignment. Returns inherited Text Style Context.
        [FIXED]: Proper Logic for Custom BG + Header Text Contrast.
        """
        def get_attr(attr_name):
            return node.get(attr_name) if node else None
        
        # 1. Base CSS Merging
        final_styles = row_styles.copy() if (not styles and row_styles) else styles.copy()
        if styles and row_styles:
             for k, v in styles.items(): final_styles[k] = v

        # -------------------------------------------------------------
        # STEP A: DETERMINE BACKGROUND (Color Selection Priority)
        # -------------------------------------------------------------
        
        # 1. Check Specific Sources in Order
        bg_self  = styles.get('background-color') or styles.get('background') or get_attr('bgcolor')
        bg_row   = row_styles.get('background-color') or row_styles.get('background') if row_styles else None
        bg_table = table_styles.get('background-color') or table_styles.get('background') if table_styles else None
        
        # 2. Pick the strongest custom background found
        # (Self > Row > Table Global)
        custom_bg_val = bg_self or bg_row or bg_table
        
        # 3. Is there actually a valid custom color?
        is_custom_bg = False
        if custom_bg_val and str(custom_bg_val).lower() not in ['transparent', 'none']:
            is_custom_bg = True

        # -------------------------------------------------------------
        # STEP B: HEADER LOGIC (THEME vs CUSTOM)
        # -------------------------------------------------------------
        
        target_bg_color = custom_bg_val
        target_text_color = None # Default Auto/Black
        
        # Extract User Defined Text Color First
        user_text = final_styles.get('color') or (row_styles.get('color') if row_styles else None)

        if is_header:
            if not is_custom_bg:
                # CASE 1: हेडर है और यूजर ने कोई बैकग्राउंड नहीं दिया -> थीम लगाओ (Blue + White)
                target_bg_color = ThemeConfig.THEME_COLORS.get('table_header_bg', 'D9E2F3') 
                if not user_text:
                    target_text_color = ThemeConfig.THEME_COLORS.get('table_header_text', 'FFFFFF')
            else:
                # CASE 2: हेडर है लेकिन यूजर ने बैकग्राउंड (#f2f2f2) दिया है -> थीम का टेक्स्ट रंग मत थोपो!
                #  Fix: Force Black for readability on custom backgrounds
                if not user_text:
                    target_text_color = '000000' # Black matches custom light BGs best

        # If it's normal TD and no custom text, keep it None (Auto/Black)
        final_text_val = user_text if user_text else target_text_color

        # -------------------------------------------------------------
        # STEP C: APPLY XML SHADING (Background)
        # -------------------------------------------------------------
        if target_bg_color:
            hex_col = ColorManager.get_hex(target_bg_color)
            if hex_col:
                XmlBuilder.set_cell_shading(cell, hex_col)

        # -------------------------------------------------------------
        # STEP D: BORDERS & WIDTH (WITH TOTAL CONTROL OVERRIDE)
        # -------------------------------------------------------------
        # 1. Width Logic (यह पहले की तरह ही रहेगा)
        width_attr = final_styles.get('width') or get_attr('width')
        if width_attr:
            if '%' in str(width_attr):
                val = UnitConverter.to_table_pct(width_attr)
                XmlBuilder.set_cell_width(cell, str(val), 'pct')
            else:
                twips = UnitConverter.to_twips(str(width_attr))
                XmlBuilder.set_cell_width(cell, str(twips), 'dxa')

        # 🚀 [TOTAL CONTROL FIX]: टकराव रोकने के लिए सुधार (Empty dict {} is VALID)
        # हम 'is not None' चेक करेंगे, न कि simple boolean check
        if explicit_border_map is not None:
             # यदि TableController ने मैप भेजा है (चाहे वह खाली {} ही क्यों न हो),
             # तो उसी का पालन करें। अपनी तरफ से दिमाग (else block) न लगाएं।
             XmlBuilder.set_cell_borders(cell, explicit_border_map)
        else:
             # Fallback: यह केवल तभी चलेगा जब TableController से कोई डेटा न आया हो (None)
             cls._apply_borders(cell, final_styles)

        # -------------------------------------------------------------
        # STEP E: RESOLVED ALIGNMENT & ROTATION (Ultimate Precision Fix)
        # -------------------------------------------------------------
        
        # 1. अलाइनमेंट रिज़ॉल्यूशन चेन
        raw_valign = (
            styles.get('vertical-align') or                
            get_attr('valign') or                         
            (row_styles.get('vertical-align') if row_styles else None) 
        )

        if getattr(AppConfig, 'DEBUG_TABLES', False):
            tag_name = node.name if node else "cell"
            logger.debug(f"   📐 [Align-Trace] Tag:{tag_name} | Raw Valign: {raw_valign}")

        # 2. कीवर्ड सिंकिंग (Default Center for all)
        if not raw_valign:
            target_valign = 'center'
        else:
            clean_val = str(raw_valign).lower().strip()
            target_valign = 'center' if clean_val in ['middle', 'center'] else clean_val

        # 3. 🚨 रोटेशन मास्टर सुधार (The Core Fix) 🚨
        writing_mode = final_styles.get('writing-mode', '').lower()
        is_rotated = 'vertical' in writing_mode or 'tb-rl' in writing_mode

        if is_rotated:
            XmlBuilder.set_cell_text_direction(cell, 'tbRl')
            # वर्ड में रोटेटेड सेल को बीच में रखने के लिए पैराग्राफ को हमेशा सेंटर होना चाहिए
            h_align = 'center'
            target_valign = 'center'
            
        # 4. XML इंजेक्ट करें
        XmlBuilder.set_cell_valign(cell, target_valign)

        # 5. हॉरिजॉन्टल अलाइनमेंट (Horizontal Alignment)
        if not is_rotated:
            h_align = final_styles.get('text-align') or \
                      get_attr('align') or \
                      (row_styles.get('text-align') if row_styles else None)

        align_enum = None
        if h_align:
            h_val = h_align.lower()
            if 'center' in h_val: align_enum = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif 'right' in h_val: align_enum = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif 'justify' in h_val: align_enum = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            else: align_enum = WD_PARAGRAPH_ALIGNMENT.LEFT

        # -------------------------------------------------------------
        # STEP E.2: RESOLVED PADDING & ROTATION CONFLICT
        # -------------------------------------------------------------
        
        # रोटेटेड सेल्स के लिए "Symmetric Minimal Padding" का उपयोग करें ताकि 
        # पैडिंग टेक्स्ट को किनारे की ओर धक्का न दे। 
        # नॉन-रोटेटेड के लिए 80/100 Twips का मानक बना रहेगा।
        
        if is_rotated:
            # रोटेटेड स्थिति में कंटेंट के केंद्र में होने की गारंटी देने के लिए '0' बेस्ट है
            default_padding = {'top': 0, 'bottom': 0, 'left': 45, 'right': 45}
        else:
            default_padding = {'top': 80, 'bottom': 80, 'left': 100, 'right': 100}

        final_margins = {}
        padding_keys = ['padding', 'padding-top', 'padding-bottom', 'padding-left', 'padding-right']
        has_user_padding = any(k in final_styles for k in padding_keys)

        # प्राथमिकता logic
        if has_user_padding:
            for side in ['top', 'bottom', 'left', 'right']:
                val = final_styles.get(f'padding-{side}') or final_styles.get('padding')
                if val is not None:
                    final_margins[side] = UnitConverter.to_twips(str(val))
                else:
                    final_margins[side] = default_padding[side]
        else:
            # अगर कोई यूजर CSS नहीं है, तो ऊपर निर्धारित बैलेंस पैडिंग लें
            final_margins = default_padding

        # Final XML call
        XmlBuilder.set_cell_margins(cell, final_margins)


        # -------------------------------------------------------------
        # STEP F: BUILD CONTEXT (Final Text Style)
        # -------------------------------------------------------------
        safe_font_styles = final_styles if final_styles else {}
        font_family = FontHandler.resolve_font_config(safe_font_styles)['ascii'] 
        if not font_family and row_styles:
             font_family = FontHandler.resolve_font_config(row_styles)['ascii']

        context = {
            'color': final_text_val, # अब यह सही 'Black' (000000) ले जाएगा
            'font_family': font_family,
            'align': align_enum,
            'bold': is_header, 
            'margin-top': '0pt',      
            'margin-bottom': '0pt', 
            'line-height': '115%',    
        }
        
        return context


    @staticmethod
    def _apply_borders(cell, styles):
        # 1. PERIMETER: ये वो साइड्स हैं जिन पर ग्लोबल 'border' लागू होना चाहिए
        PERIMETER = ['top', 'bottom', 'left', 'right']
        # 2. DIAGONAL: ये केवल तभी लगनी चाहिए जब HTML में अलग से पूछी जाएं
        DIAGONALS = ['tl2br', 'tr2bl']
        
        border_map = {}
        global_border = styles.get('border')
        
        # बाहरी ४ दिशाओं के लिए (इन्हेरिटेंस सपोर्ट के साथ)
        for side in PERIMETER:
            spec = styles.get(f'border-{side}') or global_border
            if spec:
                parsed = BorderParser.parse(spec)
                if parsed: border_map[side] = parsed

        # विकर्ण रेखाओं के लिए (कोई इन्हेरिटेंस नहीं, केवल डायरेक्ट CSS)
        for diag in DIAGONALS:
            spec = styles.get(f'border-{diag}')
            if spec:
                parsed = BorderParser.parse(spec)
                if parsed: border_map[diag] = parsed
        
        if border_map:
            XmlBuilder.set_cell_borders(cell, border_map)
            

    @classmethod
    def _apply_ghost_styling(cls, ghost_cell, master_data, row_styles):
        """Used for rowspan 'continue' cells to maintain visual border/bg integrity."""
        # Usually minimal: just shading.
        # Getting Styles from master TAG (Slow but accurate)
        if not master_data or 'tag' not in master_data: return
        
        style_str = master_data['tag'].get('style', '')
        styles = CssParser.parse(style_str)
        
        # Apply Shading only (Borders on merged continuation usually confusing in Word)
        bg = styles.get('background-color') or styles.get('background')
        if not bg and row_styles:
            bg = row_styles.get('background-color')
            
        if bg and bg != 'transparent':
            hex_val = ColorManager.get_hex(bg)
            XmlBuilder.set_cell_shading(ghost_cell, hex_val)

        cls._apply_borders(ghost_cell, styles)


    @classmethod
    def _apply_ghost_styling_shading_only(cls, ghost_cell, master_data, row_styles):
        """
        Only applies background color from Master. Borders are skipped 
        (because they are handled by Matrix Calculation).
        """
        if not master_data or 'tag' not in master_data: return
        
        style_str = master_data['tag'].get('style', '')
        styles = CssParser.parse(style_str)
        
        bg = styles.get('background-color') or styles.get('background')
        if not bg and row_styles:
            bg = row_styles.get('background-color')
            
        if bg and bg != 'transparent':
            hex_val = ColorManager.get_hex(bg)
            XmlBuilder.set_cell_shading(ghost_cell, hex_val)


    @staticmethod
    def _cleanup_trailing_space(cell):
        """
        Word adds an empty paragraph to every new cell. 
        If we added content, this empty paragraph pushes margins.
        """
        paragraphs = cell.paragraphs
        if len(paragraphs) > 1:
            last_p = paragraphs[-1]
            if not last_p.text.strip() and not last_p.runs:
                # Be careful: A cell MUST contain at least one paragraph.
                # So we modify its spacing to 0 instead of deleting,
                # unless there's preceding content.
                p_element = last_p._element
                if p_element.getparent() is not None:
                    p_element.getparent().remove(p_element)