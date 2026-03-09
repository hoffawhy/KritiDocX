"""
ROUTER CORE (The Grand Master / Traffic Controller)
---------------------------------------------------
Responsibility:
Central Dispatcher that connects parsed HTML nodes to Domain Objects.

Core Loop:
1. Receive Node from Parser.
2. Determine Content Type (Text vs Tag).
3. Merge Styles (Parent Context + Current Inline Styles).
4. Select Controller (Text, Media, Table, Form, List).
5. Execute & Recurse (Pass control down to children).

Architecture:
Uses the Strategy Pattern implicitly by selecting handlers from a map.
Manages the "Cursor" (Container) where content is injected.
"""

import bs4
import logging
import re
from docx.shared import Twips, Pt, Inches
# --- BASICS ---
from kritidocx.config.settings import AppConfig
from kritidocx.basics.css_parser import CssParser
from kritidocx.basics.color_manager import ColorManager
from kritidocx.basics.unit_converter import UnitConverter

# --- OBJECTS ---
# 1. Text
from kritidocx.objects.text import (
    ParagraphManager, RunManager, HeadingManager, 
    HyperlinkManager, BreakManager
)
# 2. Layout (used for hr, sections)
from kritidocx.objects.layout import ColumnManager # Optional helper
# 3. Media
from kritidocx.xml_factory.xml_builder import XmlBuilder # Direct low-level calls

from kritidocx.utils.style_filter import StyleFilter

# Loggers
logger = logging.getLogger("MyDocX_Router")

class Router:
    """
    The orchestrator traversing the DOM tree.
    Has access to all controllers via the Driver.
    """

    def __init__(self, doc_driver):
        """
        :param doc_driver: Instance of src.core.docx_driver.DocxDriver
        """
        self.driver = doc_driver

        # A. Handlers Map
        # Routes tag names to specific internal methods
        self._route_map = {
            # Block Level
            'p': self._handle_paragraph,
            'div': self._handle_div,
            'pre': self._handle_div,
            'article': self._handle_div, 'section': self._handle_div,
            
            # Headings
            'h1': self._handle_heading, 'h2': self._handle_heading,
            'h3': self._handle_heading, 'h4': self._handle_heading,
            'h5': self._handle_heading, 'h6': self._handle_heading,

            # Inline Level (Wrapped in Para logic mostly, but explicit handling supported)
            'span': self._handle_inline_wrapper,
            'strong': self._handle_inline_wrapper, 'b': self._handle_inline_wrapper,
            'em': self._handle_inline_wrapper, 'i': self._handle_inline_wrapper,
            'u': self._handle_inline_wrapper, 'small': self._handle_inline_wrapper,
            'strike': self._handle_inline_wrapper, 'del': self._handle_inline_wrapper,
            'sub': self._handle_inline_wrapper, 'sup': self._handle_inline_wrapper,
            'mark': self._handle_inline_wrapper,

            # Hyperlinks
            'a': self._handle_hyperlink,

            # Structure
            'br': self._handle_break,
            'hr': self._handle_horizontal_rule,
            
            # Complex Components
            'header': self._handle_header,
            'footer': self._handle_footer,
            'table': self._handle_table,
            'img': self._handle_image,
            'ul': self._handle_list,
            'ol': self._handle_list,

            # Specialized
            'math': self._handle_math,
            'input': self._handle_form,
            'select': self._handle_form,
            'textarea': self._handle_form,
            
            
        }

    #======================================================================
    # [UPDATED SECTION] COMPLEX HEADER & FOOTER HANDLERS
    # =========================================================================

    def _handle_header(self, node, container, context):
        """
        Processes <header>. Now supports Rich Content (Tables, Images, Logos).
        Uses 'HeaderFooterManager' to get the writable object, then recurses.
        """
        # 1. डीबग लॉगिंग (Logging Start)
        if getattr(AppConfig, 'DEBUG', False):
            logger.debug(f"\n🔍 [ROUTER] Processing Rich <HEADER> tag")
            logger.debug(f"   ➤ Style Context: {context}")

        # 2. सही सेक्शन और हैडर ऑब्जेक्ट प्राप्त करें
        # (हम पिछले सेक्शन को ही टारगेट करते हैं)
        section = self.driver.doc.sections[-1]
        
        from kritidocx.objects.layout.header_footer import HeaderFooterManager
        
        # यह पता करें कि यह 'First Page Header' है या 'Default' (HTML attribute 'data-type' से)
        h_type = node.get('data-type', 'default')
        is_first = (h_type == 'first')
        
        # हैडर का कंटेनर (Word Wrapper) लाएं
        header_container = HeaderFooterManager.get_active_header(section, is_first_page=is_first)

        # [UPDATE START]: Flex Layout Check for Headers
        # अगर CSS में 'flex' है, तो इसे Grid System में बदलें,
        # नहीं तो साधारण पैराग्राफ टेक्स्ट चिपका (Mashed) देगा।
        display_mode = context.get('display', '').lower()
        if 'flex' in display_mode or 'grid' in display_mode:
            
            # --- Auto-alignment fix for 'space-between' ---
            # अगर 'space-between' है, तो बच्चों पर अलाइनमेंट फोर्स करें 
            # (पहला लेफ्ट, आखिरी राइट) ताकि टेबल सेल में सही दिखें
            justify = context.get('justify-content', '')
            children = [c for c in node.children if c.name] # Only Tags
            
            if 'space-between' in justify and len(children) >= 2:
                # First child left
                first_style = children[0].get('style', '') + "; text-align: left;"
                children[0]['style'] = first_style
                # Last child right
                last_style = children[-1].get('style', '') + "; text-align: right;"
                children[-1]['style'] = last_style

            # Flex Grid Engine Call
            if self.driver.table_ctrl.create_flex_layout_grid(
                node=node,
                container=header_container,
                context=context,
                router_callback=self.process_node
            ):
                return # Table built successfully, stop standard recursion
    

        # 3. रिकर्शन (Recursion) - Standard Fallback
        self._recurse_children(node, header_container, context) 
        # 4. लॉगिंग समाप्ति (Logging End)
        if getattr(AppConfig, 'DEBUG', False):
            logger.info(f"   ✅ [ROUTER] Header Injection Complete (Layout Elements Processed).")

    def _handle_footer(self, node, container, context):
        """
        Processes <footer>. 
        [UPDATED FIX]: Iterates all sections to support mixed Portrait/Landscape layouts.
        """
        if getattr(AppConfig, 'DEBUG', False):
            logger.debug(f"\n🔍 [ROUTER] Processing Global <FOOTER>")

        # हम फूटर को "हर सेक्शन" में अलग से इंजेक्ट करेंगे
        # ताकि अगर सेक्शन A4 है तो A4 की टेबल बने, अगर Landscape है तो चौड़ी टेबल बने।
        
        doc_sections = self.driver.doc.sections
        
        # 1. लूप चलाएं सभी सेक्शन्स पर
        for idx, section in enumerate(doc_sections):
            
            # Smart Check: यदि लिंक 'True' है और यह पहला सेक्शन नहीं है, 
            # तो इसे 'False' करें ताकि हम नया टेबल डाल सकें (वर्ना पुराना ही रिपीट होगा)
            # लेकिन अगर यूजर "Same as previous" चाहता है, तो हमें कोड बदलना होगा। 
            # फिलहाल "Responsive Fix" के लिए अनलिंक करना सबसे सुरक्षित है।
            if idx > 0:
                section.footer.is_linked_to_previous = False

            # कंटेनर प्राप्त करें
            from kritidocx.objects.layout.header_footer import HeaderFooterManager
            footer_container = HeaderFooterManager.get_active_footer(section, is_first_page=False)

        display_mode = context.get('display', '').lower()
        if 'flex' in display_mode or 'grid' in display_mode:
            
            # --- Auto-alignment fix for 'space-between' ---
            # अगर 'space-between' है, तो बच्चों पर अलाइनमेंट फोर्स करें 
            # (पहला लेफ्ट, आखिरी राइट) ताकि टेबल सेल में सही दिखें
            justify = context.get('justify-content', '')
            children = [c for c in node.children if c.name] # Only Tags
            
            if 'space-between' in justify and len(children) >= 2:
                # First child left
                first_style = children[0].get('style', '') + "; text-align: left;"
                children[0]['style'] = first_style
                # Last child right
                last_style = children[-1].get('style', '') + "; text-align: right;"
                children[-1]['style'] = last_style

            # Flex Grid Engine Call
            if self.driver.table_ctrl.create_flex_layout_grid(
                node=node,
                container=footer_container,
                context=context,
                router_callback=self.process_node
            ):
                return # Table built successfully, stop standard recursion

            # [RECURSION RESTART]: हर बार राउटर कॉल करें
            # ध्यान दें: _recurse_children iterator consume कर लेता है, 
            # इसलिए हमें नोड की बच्चों की लिस्ट एक बार बना लेनी चाहिए
            # (BeautifulSoup ऑब्जेक्ट reuse हो सकता है, iterator नहीं)
            
            # बेहतर तरीका: self._recurse_children को कॉल करने के बजाय, 
            # हम सीधे बच्चों की सूची (list) पर लूप चलाएंगे ताकि iterator खत्म न हो।
            
            if hasattr(node, 'children'):
                import bs4
                children_list = [c for c in node.children] # लिस्ट में कन्वर्ट करें
                
                # Context को क्लीन करें
                from kritidocx.utils.style_filter import StyleFilter
                clean_context = StyleFilter.get_clean_child_context(context)

                # प्रत्येक बच्चे को प्रोसेस करें
                for child in children_list:
                    # हर सेक्शन के लिए फिर से टेबल जनरेट होगा (नयी Width के साथ)
                    self.process_node(child, footer_container, clean_context)

            # पेज नंबर जोड़ना
            if node.get('data-auto-page-number') == 'true':
                 HeaderFooterManager.add_page_numbers_to_container(footer_container, context)

        if getattr(AppConfig, 'DEBUG', False):
            logger.debug(f"   ✅ [ROUTER] Footer applied to {len(doc_sections)} sections adaptively.")
            
    # =========================================================================
    # 🚀 MASTER PROCESS METHOD
    # =========================================================================
    def process_node(self, node, container=None, parent_context=None):
        """
        Recursively processes a BeautifulSoup Node.
        [MASTER UPDATE]: Includes Style Leakage Protection + Ghost Paragraph Killer.
        """
        # 1. संदर्भ (Context) इनिशियलाइज़ेशन
        if parent_context is None: parent_context = {}
        
        # =========================================================================
        # [MASTER GATE]: Parent Style Leakage Prevention
        # यह सुनिश्चित करता है कि कंटेनर का बॉर्डर/बैकग्राउंड किसी भी बच्चे (Text or Tag) पर न लगे।
        # =========================================================================
        from kritidocx.utils.style_filter import StyleFilter
        inherited_safe = StyleFilter.get_clean_child_context(parent_context)
        # =========================================================================

        if not node: return

        # Default Container = Document Body
        # [CRITICAL FIX]: 'is not None' का उपयोग करें। 
        # अन्यथा खाली Textbox XML Node को Python 'False' मान लेता है और Main Doc पर स्विच हो जाता है।
        current_container = container if container is not None else self.driver.doc
        
        # --- CASE 1: PLAIN TEXT NODE ---
        if isinstance(node, bs4.element.NavigableString):
            # [PRESERVED]: GHOST PARAGRAPH KILLER
            raw_text = str(node)
            if not raw_text.strip(" \n\r\t"):
                return 

            # [FIXED]: अब यहाँ 'inherited_safe' का उपयोग होगा ताकि टेक्स्ट पर बॉर्डर न आए
            self._handle_text_node(node, current_container, inherited_safe)
            return

        if not isinstance(node, bs4.element.Tag): return

        # --- CASE 2: TAG NODES (PREPARE CONTEXT) ---
        tag = node.name.lower()
        
        # वर्तमान नोड की अपनी इनलाइन स्टाइल लें
        current_style_str = node.get('style', '')
        css_data = CssParser.parse(current_style_str)
        
        # [FIXED]: New Context = साफ़ किया हुआ पैरेंट + नोड का अपना नया स्टाइल
        new_context = inherited_safe.copy()
        new_context.update(css_data)
        
        # 🟢 [FIX]: Explicit Alignment Bridge
        # यदि पैरेंट कंटेनर (जैसे .header) सेंटर अलाइन्ड है, तो इसे इनहेरिट करें
        if parent_context.get('text-align') == 'center' and 'text-align' not in css_data:
            new_context['text-align'] = 'center'
            new_context['align'] = 'center'

        
        # [PRESERVED]: Map specific HTML attributes
        if node.get('color'): new_context['color'] = node.get('color')
        
        # --- DISPATCH ---
        handler = self._route_map.get(tag)
        if handler:
            # Execute specific logic (Heading, Table, List etc.)
            handler(node, current_container, new_context)
        else:
            # Fallback: Generic wrapper
            self._recurse_children(node, current_container, new_context)
            
    # =========================================================================
    # 🧱 BLOCK HANDLERS
    # =========================================================================

    def _handle_paragraph(self, node, container, context):
        """Standard <p>."""
        # 1. Resolve Container
        target = self._ensure_block_container(container)
        
        para = None
        
        # 2. Create Paragraph Object
        if hasattr(target, 'add_paragraph'):
            # CASE A: Standard Document/Cell
            para = target.add_paragraph()
            
        elif hasattr(target, 'append'):
            # CASE B: Raw XML Element (Textbox Content) - [CRITICAL FIX]
            # Floating boxes के लिए पैराग्राफ को मैन्युअली बनाना और चिपकाना पड़ता है
            from kritidocx.xml_factory.xml_builder import XmlBuilder
            from docx.text.paragraph import Paragraph
            
            # Temporary creation strategy to ensure styling works
            temp_p = self.driver.doc.add_paragraph()
            p_node = temp_p._element
            
            # Detach and Move
            if p_node.getparent() is not None:
                p_node.getparent().remove(p_node)
            
            target.append(p_node)
            para = temp_p  # Wrapper stays valid pointing to new location
            
        # यदि पैरा नहीं बना, तो वापस लौटें
        if not para: return

        # 3. Apply Formatting
        ParagraphManager.apply_formatting(para, context)
        
        # 4. Fill Content
        self._recurse_children(node, para, context)

    def _handle_heading(self, node, container, context):
        """Headings with Dual-Sided Spacer cleanup."""
        # 1. TOP SPACER (पुराना वाला)
        m_top = context.get('margin-top') or context.get('margin_top')
        if m_top and 'border-left' in context:
            if hasattr(container, 'add_paragraph'):
                spacer_t = container.add_paragraph()
                spacer_t.paragraph_format.space_after = Twips(UnitConverter.to_twips(str(m_top)))

        # मुख्य हेडिंग बनाने का कोड (जैसा पहले था)
        try:
            level = int(node.name[1])
        except: level = 1
        
        # क्लिप लॉजिक: यहाँ margin को हटा रहे हैं ताकि बॉर्डर न खींचे
        h_context = context.copy()
        m_bot_raw = h_context.get('margin-bottom') or h_context.get('margin_bottom') # वैल्यू सेव करें
        h_context['margin-top'] = '0pt'
        h_context['margin-bottom'] = '0pt'
        
        para = HeadingManager.add_heading(container, level, "", style_data=h_context)
        self._recurse_children(node, para, h_context)

        # --- [NEW FIX START]: BOTTOM SPACER ---
        # यदि हेडिंग में लाल लाइन है, तो उसके नीचे थोड़ा गैप (margin-bottom) अलग से जोड़ें
        # डिफ़ॉल्ट 10px-15px का गैप रखें अगर CSS में नहीं है
        m_bot = m_bot_raw if m_bot_raw else "12pt" 
        
        if 'border-left' in context or 'border-right' in context:
            if hasattr(container, 'add_paragraph'):
                spacer_b = container.add_paragraph()
                # यह अगला पैराग्राफ है जो खाली रहेगा, और 'Space After' के जरिए नीचे गैप बनाएगा
                spacer_b.paragraph_format.space_after = Twips(UnitConverter.to_twips(str(m_bot)))
                
                # इसकी लाइन हाइट भी कम करें ताकि यह खुद ज्यादा जगह न ले
                from docx.shared import Pt
                spacer_b.paragraph_format.line_spacing = Pt(1) 
        # --- [NEW FIX END] ---

    def _handle_div(self, node, container, context):
        """
        DIVs handling with Floating support + Box Container grouping fix.
        Updated: Includes Vertical Center logic & Empty Page Fix.
        """
        
                # [NEW LOGIC START]: Class-based Redirection
        # HTML 'page-header' class को असली <header> लॉजिक पर भेजें
        classes = node.get('class', [])
        
        if 'page-header' in classes:
            self._handle_header(node, container, context)
            return

        if 'page-footer' in classes:
            self._handle_footer(node, container, context)
            return

        
        style = context
        
        # 🟢 [UPDATE START: Inline-Block "Fake Checkbox" Fix] 🟢
        # समस्या: <div style="display: inline-block; width:15px; border:1px...">
        # यह "Box Logic" को ट्रिगर करता है जो एक Table बनाता है, जिससे लाइन टूट जाती है।
        # समाधान: इसे Checkbox Handler पर भेज दें।
        
        display_mode = style.get('display', '').lower()
        width_raw = style.get('width')
        is_inline_block = (display_mode == 'inline-block')
        
        # छोटा बॉक्स है? (30px/30pt से कम) - यानी चेकबॉक्स जैसा है
        is_small_box = False
        if width_raw:
            try:
                # 30 * 15 (Twips/px factor) = 450 Approx cutoff
                val_twips = UnitConverter.to_twips(str(width_raw))
                if 0 < val_twips < 600: 
                    is_small_box = True
            except: pass

        if is_inline_block and is_small_box and 'border' in str(style):
            if getattr(AppConfig, 'DEBUG_FORMS', False):
                logger.debug(f"   ☑️ [Router] Mapping Div-Box to Checkbox SDT")
                
            # इसे सीधा चेकबॉक्स हैंडलर को दें (पैराग्राफ कॉन्टेक्स्ट सुरक्षित करें)
            target_para = self._ensure_paragraph_context(container, context)
            
            # FormController से चेकबॉक्स बनाएं (Fake node pass करके)
            # Checked=False डिफ़ॉल्ट है
            from kritidocx.objects.form.checkbox_handler import CheckboxHandler
            CheckboxHandler.add_checkbox(node, target_para, context)
            return
        # 🟢 [UPDATE END] 🟢


        
        
        # 1.1 =========================================================
        # 🚀 [INTERCEPTOR START]: Floating & Rotation Guard
        # =========================================================
        # लक्ष्य: क्या इस DIV को सामान्य ब्लॉक की तरह रेंडर करना चाहिए या ग्राफिक (Shape) की तरह?
        # ट्रिगर्स: 
        # A. Position: Absolute/Fixed
        # B. Float: Left/Right
        # C. Rotation: transform: rotate(...) [यह नया एडिशन है]
        
        pos_mode = style.get('position', '').lower()
        float_mode = style.get('float', '').lower()
        # CssEngine 'rotation_deg' key में numeric value डालता है (e.g. 45.0)
        has_rotation = (style.get('rotation_deg') is not None) 

        is_floating_graphic = (
            (pos_mode in ['absolute', 'fixed']) or 
            (float_mode in ['left', 'right']) or 
            has_rotation
        )
        
        if is_floating_graphic:
            # 1. Add Textbox
            textbox_content_xml_node = self.driver.media_ctrl.add_floating_textbox(
                node_info={'style_dict': style},
                container=self._ensure_paragraph_context(container)
            )
            
            if textbox_content_xml_node is not None:
                # 3. RECURSION (With Context Cleaning)
                
                # [FIX]: Context Leaking Solution
                # हम सीधे 'context' (style) नहीं भेजेंगे। हम उसे Filter करेंगे।
                # यह Textbox की Border/Background को Text Run पर लगने से रोकेगा।
                from kritidocx.utils.style_filter import StyleFilter
                
                child_safe_style = StyleFilter.get_clean_child_context(style)
                
                self._recurse_children(node, textbox_content_xml_node, child_safe_style)
            
            return
        
        # [INTERCEPTOR END]
        # =========================================================
        # =========================================================
        # 🟢 [REFACTORED STEP 3] FLEX/GRID HANDLER 
        # =========================================================
        # लक्ष्य: अगर DIV एक 'Flex/Grid' है, तो सीधे Object Level Controller को कॉल करें।
        display_mode = style.get('display', '').lower()
        
        if display_mode in ['flex', 'grid', 'inline-flex', 'inline-grid']:
            # हमने जो Step 2 में नया शुद्ध Logic वाला फंक्शन बनाया है, उसे यहाँ से ट्रिगर करें।
            # हम container (Document/Cell), node (Div), context (Styles) और
            # process_node (Recursion Callback) भेज रहे हैं।
            if self.driver.table_ctrl.create_flex_layout_grid(
                node=node, 
                container=container, 
                context=context, 
                router_callback=self.process_node
            ):
                # यदि Flex Grid सफलतापूर्वक बन गया है, तो यहीं से बाहर निकलें (return) 
                # ताकि नीचे का 'Default Box' या 'Recursion' दोबारा न चले (Duplication रोकेगा)।
                return 

   # =========================================================
        # 1.2 LAYOUT TRIGGERS DETECTION (लेआउट ट्रिगर्स की पहचान)
        # =========================================================
        
        has_break = style.get('page-break-before') == 'always' or style.get('break-before') == 'page'
        has_columns = style.get('column-count') is not None
        has_page_size = style.get('size') is not None
        
        # ✅ NEW LOGIC: चेक करें कि क्या यूजर वर्टिकल सेंटर चाहता है
        # 1. Flexbox तरीका: display: flex + align-items: center
        # 2. Classic तरीका: vertical-align: middle या center
        is_flex_center = (style.get('display') == 'flex' and style.get('align-items') == 'center')
        is_valign_center = (style.get('vertical-align') in ['middle', 'center'])

        # Layout Logic ट्रिगर करें (Page Break, Size, Columns या Vertical Align)
        if has_break or has_columns or has_page_size:
            
            doc = self.driver.doc
            
            # --- DIAGNOSTIC LOGIC: (क्या यह डॉक्यूमेंट का पहला पन्ना है?) ---
            # अगर पहले पन्ने पर कंटेंट नहीं है, तो नया ब्रेक लगाने की जगह हम उसी पन्ने को रीसायकल (Reuse) करेंगे।
            para_count = len(doc.paragraphs)
            has_tables = len(doc.tables) > 0
            
            # सुरक्षित टेक्स्ट चेक (ताकि खाली लिस्ट पर क्रैश न हो)
            last_para_text = doc.paragraphs[-1].text if para_count > 0 else ""
            
            is_start_of_doc = (
                len(doc.sections) == 1 and 
                not has_tables and 
                (para_count == 0 or (para_count <= 2 and not last_para_text.strip()))
            )

            # Debug Output
            if getattr(AppConfig, 'DEBUG', False):
                status_msg = "REUSE (Fix Page 1)" if is_start_of_doc else "NEW PAGE (Break)"
                logger.debug(f"   🛑 [Layout] P:{para_count} | Txt:'{last_para_text[:10]}' | Dec:{status_msg}")

            # --- SECTION CREATION LOGIC ---
            if is_start_of_doc:
                # यदि शुरुआत है, तो मौजूदा सेक्शन का उपयोग करें
                new_section = doc.sections[-1]
            else:
                # अन्यथा नया पन्ना/सेक्शन जोड़ें
                new_section = self.driver.section_mgr.add_section_break('next_page')
                
                # 🛡️ [FIX]: Break Linkage from Previous Section
                # जब हम लेआउट बदलते हैं (Portrait -> Landscape), तो पुराना Header/Footer
                # फिट नहीं बैठता। इसलिए लिंक तोड़ दें।
                if has_break or has_columns or has_page_size:
                    new_section.header.is_linked_to_previous = False
                    new_section.footer.is_linked_to_previous = False
                    
                    # यह सुनिश्चित करता है कि पिछले सेक्शन का टेबल ग्रिड (Portrait)
                    # नए सेक्शन (Landscape) पर जबरदस्ती लागू न हो।
                    if getattr(AppConfig, 'DEBUG', False):
                        logger.debug(f"   ⚓ [Layout] Unlinked Header/Footer for New Section Geometry.")
                        
                        
            # =========================================================
            # A. ✅ VERTICAL ALIGNMENT FIX (नया कोड यहाँ जोड़ा गया है)
            # =========================================================
            if is_flex_center or is_valign_center:
                self.driver.section_mgr.set_vertical_alignment('center', section=new_section)
                if getattr(AppConfig, 'DEBUG', False):
                    logger.debug(f"   📐 [Layout] Applied Vertical Center to Page.")

            # =========================================================
            # B. COLUMN APPLICATION (पुराना कोड सुरक्षित)
            # =========================================================
            from kritidocx.objects.layout.column_manager import ColumnManager
            if has_columns:
                count = style.get('column-count')
                gap = style.get('column-gap')
                rule = (style.get('column-rule-width', '0') != '0' or 'solid' in style.get('column-rule', ''))
                ColumnManager.apply_columns(new_section, num_columns=count, spacing=gap, separator=rule)
                
            # =========================================================
            # C. PAGE SIZE & ORIENTATION (पुराना कोड सुरक्षित)
            # =========================================================
            if has_page_size:
                # Lazy Import to prevent Circular issues
                from kritidocx.objects.layout.page_setup import PageSetup
                PageSetup.set_custom_size_from_css(new_section, style)
            
            # =========================================================
            # D. MARGINS (भविष्य के लिए सुरक्षित रखा गया है)
            # =========================================================
            if any(k in style for k in ['margin-left', 'margin-top', 'margin', 'padding-left', 'padding-top']):
                from kritidocx.objects.layout.margin_manager import MarginManager
                # MarginManager needs to act on the 'new_section' or the section where the cursor currently is
                target_section = new_section if 'new_section' in locals() else doc.sections[-1]
                
                MarginManager(self.driver.doc).apply_margins(target_section, style_data=style)
                if getattr(AppConfig, 'DEBUG', False):
                    logger.debug(f"   📏 [Layout] Applied Margins.")


            # [CRITICAL]: रिकर्शन कॉल करें लेकिन नया Container न बदलें
            # हम इसे सिर्फ "Layout Instruction" मान रहे हैं।
            self._recurse_children(node, container, context)
            return

        # : Floating/Absolute logic (Textbox Shape)
        pos_mode = style.get('position', '').lower()
        float_mode = style.get('float', '').lower()
        is_floating = (pos_mode in ['absolute', 'fixed']) or (float_mode in ['left', 'right'])
        
        if is_floating:
            # 1. Textbox (Shape) बनाएँ और उसके अंदर का खाली कंटेनर (xml node) प्राप्त करें
            textbox_content_node = self.driver.media_ctrl.add_floating_textbox(
                node_info={'style_dict': style},
                container=self._ensure_paragraph_context(container) 
            )
            
            # 2. ✅ CRITICAL FIX: केवल यदि बॉक्स बना, तो रिकर्शन उसमें करें और RETURN करें।
            # इससे टेक्स्ट Main Body में जाने से बच जाएगा।
            if textbox_content_node is not None:
                self._recurse_children(node, textbox_content_node, context)
            return


        # 🟢 [NEW IMPROVEMENT]: BOX MODEL COHESION (Table Box)
        # अगर फ्लोटिंग नहीं है, लेकिन स्टाइल में बॉर्डर या बैकग्राउंड है, 
        # तो सबको एक ही बॉक्स में रखने के लिए 1x1 Table का उपयोग करें।
        has_border = any(k for k in style if 'border' in k and style[k] not in ['none', '0', '0px'])
        has_background = any(k for k in style if 'background' in k and style[k] != 'transparent')

        if has_border or has_background:
            # TableController को कॉल करें जो एक ही सेल के अंदर पूरे कंटेंट को सुरक्षित रखेगा
            # नोट: ड्राइवर में table_ctrl.create_box_container होना अनिवार्य है।
            inner_box_container = self.driver.table_ctrl.create_box_container(style, container)
            
            # अब साफ किए गए (inner_context) के साथ बच्चों को रेंडर करें
            self._recurse_children(node, inner_box_container, context)
            return


        # Standard Wrapper behavior (नो बॉर्डर, नो बैकग्राउंड)
        # सीधा रिकर्शन बिना एक्स्ट्रा बॉक्स बनाए।
        self._recurse_children(node, container, context)

    # =========================================================================
    # 🏃 INLINE HANDLERS
    # =========================================================================

    def _handle_inline_wrapper(self, node, container, context):
        """Spans, Bold, Italic wrappers."""
        # 1. Update Context Flags based on Tag
        # Create a mutation of context for children
        mutated_context = context.copy()
        
        # 🟢 [HIGHLIGHT FIX]: Persist Background for Inline Elements
        # StyleFilter सामान्यतः 'background-color' को हटा देता है (Container safety).
        # Span के लिए हम इसे 'highlight' की में सुरक्षित कर लेंगे ताकि RunManager इसे देख सके।
        bg = mutated_context.get('background-color') or mutated_context.get('background')
        if bg and bg != 'transparent':
            mutated_context['highlight'] = bg

        t = node.name.lower()
        if t in ['b', 'strong']: mutated_context['bold'] = True
        if t in ['i', 'em']: mutated_context['italic'] = True
        if t in ['u', 'ins']: mutated_context['underline'] = True
        if t in ['strike', 'del', 's']: mutated_context['strike'] = True
        if t == 'sub': mutated_context['sub'] = True
        if t == 'sup': mutated_context['sup'] = True
        if t == 'mark': mutated_context['highlight'] = 'yellow' # Default highlight
        if t == 'code': 
            mutated_context['font_family'] = 'Courier New'
            mutated_context['highlight'] = 'light_gray' # Optional visual cue

        # 2. Check container
        # Inline elements must live in a Paragraph.
        target_para = self._ensure_paragraph_context(container)
        
        # 3. Recurse
        self._recurse_children(node, target_para, mutated_context)


    def _handle_hyperlink(self, node, container, context):
        """Anchor tags <a>."""
        target_para = self._ensure_paragraph_context(container)
        href = node.get('href')
        text = node.get_text()
        
        # HyperlinkManager handles creation
        HyperlinkManager.add_hyperlink(
            target_para, 
            href, 
            text_content=text, 
            style_data=context
        )
        # Note: We consume text via get_text(). If <a> has image children, logic needs update.
        # Currently Text-Only links supported.

    def _handle_text_node(self, nav_string, container, context):
        """
        [FINAL FIX]: Variable Name Consistency ('text') + Math Regex + Preformatted.
        """
        import re # Safety Import
        
        # 1. Capture Raw String
        text = str(nav_string)
        if not text: return

        # 2. Detect Preformatted Context
        parent_tag = nav_string.parent.name.lower() if nav_string.parent else ""
        is_preformatted = parent_tag in ['pre', 'code', 'textarea'] or \
                          'Courier New' in str(context.get('font_family', '')) or \
                          context.get('white-space') in ['pre', 'pre-wrap']

        # 3. Ensure Target Paragraph Exists
        target_para = self._ensure_paragraph_context(container, context)

        # 4. Processing Logic
        if is_preformatted:
            # === CASE A: CODE BLOCKS (Preserve formatting) ===
            from kritidocx.objects.text.run_manager import RunManager
            from kritidocx.objects.text.break_manager import BreakManager
            
            # सिर्फ Carriage Return को हटाएं, Tabs को स्पेस में बदलें
            safe_text = text.replace('\r', '').replace('\t', '    ')
            # न्यूलाइन पर स्प्लिट करें
            lines = safe_text.split('\n')
            
            for i, line_content in enumerate(lines):
                if line_content:
                    RunManager.create_run(target_para, line_content, context)
                # आखिरी लाइन को छोड़कर हर लाइन के बाद <br> जोड़ें
                if i < len(lines) - 1:
                    BreakManager.add_break(target_para, 'line')
        else:
            # === CASE B: NORMAL TEXT FLOW ===
            
            # A. Whitespace Cleaning
            text = text.replace('\r', '').replace('\n', ' ').replace('\t', ' ')
            text = re.sub(r'\s+', ' ', text)
            if not text.strip(): return

            # B. [CORRECT ORDER]: Math Detection FIRST ($$...$$)
            # यह पहले Math को तोड़ेगा, ताकि उसके अंदर के {} Form Logic में न फँसें
            math_pattern = r'(\$\$.*?\$\$)|(\$.*?\$)'
            parts = re.split(math_pattern, text)

            from kritidocx.objects.text.run_manager import RunManager
            
            for part in parts:
                if not part: continue
                
                # Check for Math Delimiters
                if (part.startswith('$$') and part.endswith('$$')) or \
                   (part.startswith('$') and part.endswith('$')):
                    
                    # Remove delimiters ($$) to get raw LaTeX (\begin{pmatrix}...)
                    clean_latex = part.strip('$').strip()
                    
                    # Send to Math Controller (Safe from Field processing)
                    self.driver.math_ctrl.process_math(
                        content=clean_latex, 
                        container=target_para, 
                        style_data=context, 
                        is_latex=True
                    )
                else:
                    # C. Field Detection (Only on non-math text)
                    # "Click here { DATE }" -> Handle Fields
                    if '{' in part and '}' in part:
                         self.driver.form_ctrl.process_field_patterns(part, target_para)
                    else:
                        # Standard Text Run
                        RunManager.create_run(target_para, part, context)
    
                    
        # [DEBUG LOG]
        if getattr(AppConfig, 'DEBUG_LISTS', False) and context.get('num_id'):
            logger.debug(f"   📝 [Router->Text] Plain Text Node in List | Context ID: {context.get('num_id')}")

        # [CRITICAL LIST FIX]: Apply List Props ONLY
        # यहाँ हम context को फ़िल्टर (Filter) कर रहे हैं।
        # यदि हम पूरा context भेजते हैं, तो बच्चे (Span) के स्टाइल्स (जैसे Border/Background)
        # गलती से पिता (Paragraph) पर लग जाते हैं। 
        # इसलिए हम केवल लिस्टिंग जानकारी पास करेंगे।
        
        if context.get('num_id'):
            list_safe_context = {
                'num_id': context.get('num_id'),
                'list_depth': context.get('list_depth'),
                # (Optional) alignment पैराग्राफ लेवल प्रॉपर्टी है, इसे रख सकते हैं
                'align': context.get('align'), 
                'text-align': context.get('text-align'),
                
                                # --- NEW ADDITION START ---
                'padding-left': context.get('padding-left'),
                'margin-left': context.get('margin-left'),
                'padding_left': context.get('padding_left'),
                'text-indent': context.get('text-indent'),
                'text_indent': context.get('text_indent')

            }
            ParagraphManager.apply_formatting(target_para, list_safe_context)
    # =========================================================================
    # 📦 COMPLEX OBJECT HANDLERS
    # =========================================================================

    def _handle_table(self, node, container, context):
        """Delegates to Table Controller with List Indent support."""
        target = self._ensure_block_container(container)
        
        # --- [NEW FIX START]: Calculate List Indent for Table ---
        indent_twips = 0
        if context.get('list_depth') is not None:
            # लिस्ट की गहराई के आधार पर इंडेंट निकालें
            # लेवल 0 = 720 Twips (0.5 inch), लेवल 1 = 1440...
            from kritidocx.objects.list.indent_math import IndentMath
            # हम IndentMath का उपयोग करेंगे ताकि गणना सटीक रहे
            indent_twips, _ = IndentMath.calculate(context['list_depth'], context)
        # --- [NEW FIX END] ---

        # अब Table Controller को इंडेंट पास करें
        self.driver.table_ctrl.process_table(
            node, 
            container=target, 
            parent_context=context,
            router_callback=self.process_node,
            indent_override=indent_twips  # <--- यह नया पैरामीटर भेजें
        )

    def _handle_image(self, node, container, context):
        """Delegates to Media Controller."""
        src = node.get('src')
        alt = node.get('alt')
        
        # Image can be inline (inside P) or block (inside Cell/Doc)
        # MediaController intelligently handles both.
        self.driver.media_ctrl.add_image(
            src, 
            container=container, # Controller determines wrapper need
            style_data=context, 
            alt_text=alt
        )

    def _handle_list(self, node, container, context):
        """Delegates to List Controller."""
        target = self._ensure_block_container(container)
        
        # Passes 'self' (Router) so list controller can iterate <li>
        self.driver.list_ctrl.process_list(
            node, 
            container=target, 
            parent_context=context, 
            router_instance=self # ⚠️ CRITICAL RECURSION LINK
        )

    def _handle_form(self, node, container, context):
        """Delegates to Form Controller (Input/Select)."""
        target = self._ensure_paragraph_context(container, context) # Context यहाँ पहले से use हो रहा है formatting के लिए
        
        # [CHANGE]: 'context' को आगे पास करें
        self.driver.form_ctrl.process_node(node, target, context)

    def _handle_math(self, node, container, context):
        """Delegate to Math Controller."""
        # Math lives in P, but Block math might need alignment
        self.driver.math_ctrl.process_math(
            content=node, 
            container=container,
            style_data=context, 
            is_latex=False # Tag is <math>
        )

    # =========================================================================
    # 🧱 STRUCTURE HANDLERS
    # =========================================================================

    def _handle_break(self, node, container, context):
        """<br> tag."""
        if hasattr(container, 'add_run'): # Inside Paragraph
            BreakManager.add_break(container, 'line')
        else:
            # At document root, means Empty Paragraph
            if hasattr(container, 'add_paragraph'):
                container.add_paragraph()

    def _handle_horizontal_rule(self, node, container, context):
        """<hr> tag."""
        target = self._ensure_block_container(container)
        p = None
        if hasattr(target, 'add_paragraph'):
            p = target.add_paragraph()
        else:
            # Fallback for XML
            pass 
        
        if p:
            # Use Paragraph border logic bottom to simulate HR
            XmlBuilder.set_paragraph_border(p, side='bottom', size=6) # 0.75pt line

    # =========================================================================
    # 🕵️ HELPER UTILITIES (Safety Nets)
    # =========================================================================

    def _recurse_children(self, node, container, context):
        """
        [UNIVERSAL FIX]: सभी टैग्स के बच्चों के लिए कॉन्टेक्स्ट को साफ करता है।
        """
        if hasattr(node, 'children'):
            # यहां 'StyleFilter' का जादू चलेगा
            # यह सुनिश्चित करेगा कि पैरेंट का बॉर्डर/बैकग्राउंड बच्चों (text) पर न जाए
            child_safe_context = StyleFilter.get_clean_child_context(context)
            
            for child in node.children:
                self.process_node(child, container, child_safe_context)


    def _ensure_paragraph_context(self, container, context=None):
        """
        [ULTIMATE FIX + LOGGING]: Merges Smart Reuse, Spacer Guard, and Textbox Logic.
        With Debugging capability to trace Flow.
        """
        # --- DEPENDENCIES ---
        from docx.text.paragraph import Paragraph
        from docx.shared import Pt
        from kritidocx.xml_factory.xml_builder import XmlBuilder
        from kritidocx.objects.text.paragraph_manager import ParagraphManager

        # 1. SETUP & DIAGNOSTICS
        debug_on = getattr(AppConfig, 'DEBUG_ROUTER_CONTEXT', False)
        
        target_para = None
        created_new = False

        if debug_on:
            # पहचानें कि Container क्या है? (Doc, Cell, या अजीबोगरीब XML)
            c_type = type(container).__name__
            if hasattr(container, 'tag'): c_type = f"RAW XML <{container.tag.split('}')[-1]}>"
            logger.debug(f"🔍 [CTX] Request Context for: {c_type}")

        # --- A. SIMPLE REUSE (अगर यह पहले से Writable है) ---
        if hasattr(container, 'add_run'):
            if debug_on: logger.debug("   ✅ reuse: Already a Paragraph/Run.")
            return container
            
        # --- B. SMART PARAGRAPH REUSE (Standard Doc/Cells) ---
        # केवल तभी चेक करें जब यह 'python-docx' का ऑब्जेक्ट हो (.paragraphs attribute exists)
        if hasattr(container, 'paragraphs') and len(container.paragraphs) > 0:
            last_p = container.paragraphs[-1]
            
            # 🔥 [SPACER TRAP DETECTION]
            is_trap_spacer = False
            pf = last_p.paragraph_format
            
            # Check 1: क्या Height Rule 'Exactly' (Enum 4) है?
            if pf.line_spacing_rule is not None:
                rule_val = getattr(pf.line_spacing_rule, 'value', pf.line_spacing_rule)
                if str(rule_val) == '4': 
                    # Check 2: क्या Height 5pt से कम है?
                    if pf.line_spacing and pf.line_spacing < Pt(5):
                        is_trap_spacer = True
            
            if not is_trap_spacer:
                if debug_on: logger.debug("   ♻️  reuse: Using Last Active Paragraph.")
                return last_p
            else:
                if debug_on: logger.debug("   🚫 skip: 1pt Spacer detected. Creating New.")

        # --- C. CREATION LOGIC (DECISION TREE) ---
        
        # Scenario 1: Standard Containers (Body, Cell)
        if hasattr(container, 'add_paragraph'):
            if debug_on: logger.debug("   ✨ create: Standard python-docx Paragraph.")
            target_para = container.add_paragraph()
            created_new = True

        # Scenario 2: Raw XML Nodes (Textboxes / w:txbxContent)
        else:
            if debug_on: 
                logger.debug("   ⚓ create: RAW XML mode (Floating Textbox detected).")
                logger.debug("      -> Creating detached Paragraph object.")

            # --- 🔥 UPDATED FIX: Move-and-Adopt Strategy ---
            # 1. Main Document पर एक असली पैराग्राफ बनाएं 
            # (ताकि styles.xml और document defaults सही से लिंक हो सकें)
            temp_para = self.driver.doc.add_paragraph()
            
            # 2. उसका XML एलिमेंट (Node) निकालें
            p_node = temp_para._element
            
            # 3. Main Document से डिटैच (हटाएं) करें
            if p_node.getparent() is not None:
                p_node.getparent().remove(p_node)
                
            # 4. नए कंटेनर (Textbox Content) में चिपकाएं
            container.append(p_node)
            
            # 5. रैपर वापस करें
            # (नोट: 'temp_para' अब नए स्थान पर पॉइंट कर रहा है, क्योंकि यह लाइव XML रैपर है)
            target_para = temp_para
            created_new = True

        # --- D. INITIAL STYLING ---
        if created_new and context:
            ParagraphManager.apply_formatting(target_para, context)

        return target_para

    def _ensure_block_container(self, container):
        """
        [CRITICAL FIX]: Ensures block elements go into Body, Cell, or Floating Box.
        """
        # Case 1: Standard python-docx Container (Doc/Cell)
        if hasattr(container, 'add_table'):
            return container

        # Case 1.5: Raw XML Container (Floating Textbox/Shape Content) - [NEW FIX]
        # Textboxes में 'add_table' नहीं होता, पर उनके पास 'append' (list method) होता है।
        if hasattr(container, 'append') and not hasattr(container, 'add_run'):
            return container

        # Case 2: Check Parent (e.g., if we are inside a Paragraph, go up)
        if hasattr(container, '_parent'):
            parent = container._parent
            if hasattr(parent, 'add_table'):
                return parent

        # Case 3: Fallback
        return self.driver.doc
    
    