"""
DOCX DRIVER MODULE (The Document Container)
-------------------------------------------
Responsibility:
1. Wraps the raw `python-docx` Document object.
2. Initializes and holds instances of all Domain Controllers (Text, Table, Media, etc.).
3. Manages file I/O (Saving, Metadata, Cleanup).
4. Handles XML part initialization quirks (e.g. forcing numbering.xml creation).

Access:
The `Router` interacts primarily with this Driver to access logic controllers.
"""

import os
import logging
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configuration
from kritidocx.config.settings import AppConfig

# Domain Objects (The Controllers we built)
from kritidocx.objects.table import TableController
from kritidocx.objects.media import MediaController
from kritidocx.objects.list import ListController
from kritidocx.objects.math import MathController
from kritidocx.objects.form import FormController
from kritidocx.objects.layout import SectionManager, HeaderFooterManager, PageSetup
import zipfile
import xml.dom.minidom

logger = logging.getLogger("MyDocX_Core")

class DocxDriver:
    """
    The central owner of the Word Document instance.
    Connects Low-level API with High-level Logic.
    """

    def __init__(self, template_path=None):
        """
        Args:
            template_path (str): Optional .docx to use as base template.
        """
        logger.info("🔌 Initializing Docx Driver...")
        
        # 1. Initialize python-docx Document
        if template_path and os.path.exists(template_path):
            self.doc = Document(template_path)
        else:
            self.doc = Document()

        # 2. XML Part Hacks (Critical for Engine Stability)
        # Force initialization of numbering part to avoid crashes in ListController
        self._ensure_numbering_part_exists()


        # [NEW ADDITION]: Activate Global Auto Update
        self._enable_field_updates()  
        
        # 3. Instantiate Controllers (Dependency Injection)
        # We pass 'self' (the driver) or 'self.doc' as needed by controllers
        
        self.table_ctrl = TableController(self)  # Table Engine
        self.media_ctrl = MediaController(self)  # Image/Shape Engine
        self.list_ctrl  = ListController(self)   # Bullet/Number Engine
        self.math_ctrl  = MathController(self)   # Equation Engine
        self.form_ctrl  = FormController()       # Interactive Forms
        # Layout Managers
        self.section_mgr = SectionManager(self.doc)
        
        # 4. Set Default Metadata
        self.set_metadata(
            title="Generated Report",
            subject="Automated via MyDocX 3.0",
            author="MyDocX Engine",
            keywords="python, html-to-docx, automation"
        )
        
        logger.debug("✅ Drivers & Controllers Ready.")

    # =========================================================================
    # 🛠️ SYSTEM UTILITIES
    # =========================================================================

    def _enable_field_updates(self):
        """
        Global Settings: Forces automatic Field update (Page Num) AND 
        Disables automatic spell-check red lines on startup.
        """
        try:
            settings = self.doc.settings.element
            
            # 1. Update Fields Force (ताकि पेज नंबर 1 से 2 हो जाए)
            update_tag = settings.find(qn('w:updateFields'))
            if update_tag is None:
                update_tag = OxmlElement('w:updateFields')
                settings.append(update_tag)
            update_tag.set(qn('w:val'), 'true')

            # ✅ 2. Hide Spelling Errors (लाल रेखाओं को छुपाएं)
            # <w:hideSpellingErrors w:val="true"/>
            hide_spell = settings.find(qn('w:hideSpellingErrors'))
            if hide_spell is None:
                hide_spell = OxmlElement('w:hideSpellingErrors')
                settings.append(hide_spell)
            hide_spell.set(qn('w:val'), 'true')

            # ✅ 3. Hide Grammar Errors
            hide_gram = settings.find(qn('w:hideGrammarErrors'))
            if hide_gram is None:
                hide_gram = OxmlElement('w:hideGrammarErrors')
                settings.append(hide_gram)
            hide_gram.set(qn('w:val'), 'true')

        except Exception:
            pass

    def _ensure_numbering_part_exists(self):
        """
        Fix: python-docx doesn't create 'numbering.xml' until a list is added.
        Our NumberingManager requires it to exist to inject styles.
        Hack: Create a dummy list item and delete it IMMEDIATELY.
        """
        try:
            # Check if part exists
            _ = self.doc.part.numbering_part
        except (NotImplementedError, AttributeError):
            # 1. Create temporary paragraph to force XML part creation
            p = self.doc.add_paragraph()
            
            # 2. 🟢 CRITICAL: Delete it instantly so document stays empty/clean
            try:
                p._element.getparent().remove(p._element)
                # Note: internal list 'doc.paragraphs' might technically hold ref until refresh,
                # but XML is clean. Next property access refreshes it.
            except:
                pass

    def set_metadata(self, title=None, author=None, subject=None, keywords=None):
        """Updates File Properties (Details tab in Windows)."""
        core_props = self.doc.core_properties
        if title: core_props.title = title
        if author: core_props.author = author
        if subject: core_props.subject = subject
        if keywords: core_props.keywords = keywords

    # =========================================================================
    # 💾 SAVING & OUTPUT
    # =========================================================================

    def save(self, filepath=None):
        """
        Saves the file to disk AND generates a Debug XML alongside it.
        """
        target = filepath or AppConfig.DEFAULT_FILENAME
        
        # 1. Prepare Directory
        folder = os.path.dirname(target)
        if folder and not os.path.exists(folder):
            try:
                os.makedirs(folder)
            except OSError as e:
                logger.error(f"Cannot create directory {folder}: {e}")
                return

        # 2. Cleanup
        self._remove_trailing_empty_paragraph()

        # 3. Write DOCX & XML
        try:
            # A. Save standard DOCX
            self.doc.save(target)
            logger.info(f"💾 Document successfully saved to: {target}")
            
            # B. [NEW FEATURE]: Generate Companion XML File
            self._dump_internal_xml(target)

        except PermissionError:
            logger.error(f"❌ PERMISSION DENIED: File '{target}' is open in Word.")
            logger.error("\n🚨 ERROR: कृपया पहले पुरानी Word फाइल (generated file) को बंद करें!\n")
        except Exception as e:
            logger.critical(f"❌ Critical Save Error: {e}")
            
            
    def _remove_trailing_empty_paragraph(self):
        """
        [FIXED LOGIC]: Removes extra empty paragraphs at the end.
        Improvement: Checks specifically for Drawing (Images) and Math equations
        before declaring a paragraph "Empty".
        """
        # 🟢 Use a loop to clean multiple ghosts (e.g. init hack + default)
        while len(self.doc.paragraphs) > 1:
            
            last_p = self.doc.paragraphs[-1]
            xml = last_p._element
            
            # 1. Text Check (Standard)
            has_text = bool(last_p.text.strip())
            
            # 2. 🛡️ ASSET DETECTION FIX (Check inside the raw XML for objects)
            # क्या पैराग्राफ में कोई ड्राइंग (इमेज) या मैथ फॉर्मूला है?
            # getchildren() पुराना हो चुका है, हम XPath या सीधे String Search का उपयोग करेंगे जो तेज़ है।
            raw_xml = str(xml.xml)
            
            has_media = ('<w:drawing' in raw_xml) or ('<w:pict' in raw_xml) # Images/Shapes
            has_math = ('<m:oMath' in raw_xml) # Equations
            
            # If ANY content exists, STOP deletion logic.
            if has_text or has_media or has_math:
                break
                
            # Table Safety Check
            prev = xml.getprevious()
            if prev is not None and prev.tag.endswith('tbl'):
                # Word Rule: Table must be followed by Paragraph.
                break

            # Delete the truly empty paragraph
            try:
                xml.getparent().remove(xml)
            except:
                break 

        # Final sanity check: If Doc became empty (rare race condition), fix it.
        if len(self.doc.paragraphs) == 0:
             self.doc.add_paragraph()
             
             
    # =========================================================================
    # 🔌 FALLBACK ACCESSORS
    # =========================================================================
    
    def get_document(self):
        return self.doc
        
    def add_paragraph_raw(self):
        """Direct access for Router defaults."""
        return self.doc.add_paragraph()
    
    # -------------------------------------------------------------
    # 🔎 DEBUGGING UTILITY: XML EXTRACTOR
    # -------------------------------------------------------------
    def _dump_internal_xml(self, docx_path):
        """
        Extracts 'word/document.xml' from the saved DOCX, formats it, 
        and saves it as '.xml' for inspection.
        Example: output.docx -> output.xml
        """
        if not getattr(AppConfig, 'DEBUG', False):
            return

        try:
            # .docx -> .xml Path
            xml_target_path = docx_path.replace('.docx', '.xml')
            
            # Unzip and Read
            with zipfile.ZipFile(docx_path, 'r') as z:
                # 'word/document.xml' वह फाइल है जहाँ सारा टेक्स्ट/कोड होता है
                xml_content = z.read('word/document.xml').decode('utf-8')
                
                # Prettify (Make readable with indentation)
                dom = xml.dom.minidom.parseString(xml_content)
                pretty_xml = dom.toprettyxml(indent="  ")
                
                # Remove annoying empty lines caused by toprettyxml (Optional polish)
                pretty_xml = "\n".join([line for line in pretty_xml.split('\n') if line.strip()])

            # Save
            with open(xml_target_path, "w", encoding="utf-8") as f:
                f.write(pretty_xml)
                
            logger.info(f"   📄 Debug XML exported to: {os.path.basename(xml_target_path)}")

        except Exception as e:
            logger.warning(f"⚠️ XML Dump Failed: {e}")