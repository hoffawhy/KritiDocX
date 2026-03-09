"""
XML BUILDER FACADE (The Main Control Board)
-------------------------------------------
जिम्मेदारी: 
सभी Objects और XML Factories के बीच की एकमात्र कड़ी (Single Interface).

Design Pattern: Facade.
Why? ताकी 'Business Logic' लेयर को कभी भी XML के नेम्सपेस या सॉर्टिंग के बारे में चिंता न करनी पड़े।

Features:
- Smart Delegation: Requests को सही स्पेशलिस्ट के पास भेजना।
- Complex Assembly: इमेज और टेक्स्टबॉक्स के नेस्टेड स्ट्रक्चर को असेंबल करना।
- Type Safety: इनपुट्स को साफ (Clean) करना।
"""

from kritidocx.config.settings import AppConfig
from kritidocx.utils.logger import logger
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Import all specialist factories
from .base import XmlBase
from .text_xml import TextXml
from .table_xml import TableXml
from .drawing_xml import DrawingXml
from .form_xml import FormXml
from .section_xml import SectionXml
from .numbering_xml import NumberingXml

class XmlBuilder(XmlBase):
    """
    Public API for XML generation.
    Static methods callable from anywhere in the codebase.
    """

    # =========================================================================
    # 1. 📝 TEXT OPERATIONS (Fonts, Colors, Spacing)
    # =========================================================================

    @staticmethod
    def set_paragraph_shading(paragraph, hex_color):
        TextXml.set_paragraph_shading(paragraph, hex_color)

    @staticmethod
    def set_paragraph_border(paragraph, **kwargs):
        """Arguments: side, size, color, style"""
        TextXml.set_paragraph_border(paragraph, **kwargs)

    @staticmethod
    def set_paragraph_indent(paragraph, left=0, right=0, first_line=0, hanging=0):
        TextXml.set_indent(paragraph, left, right, first_line, hanging)

    @staticmethod
    def set_paragraph_spacing(paragraph, **kwargs):
        """Args: before, after, line, line_rule"""
        TextXml.set_spacing(paragraph, **kwargs)

    @staticmethod
    def set_paragraph_numbering(paragraph, level, num_id):
        # FIX: Explicitly map arguments to match TextXml definition (num_id first, then level)
        TextXml.set_numbering(paragraph, num_id=num_id, level_id=level)
 
    @staticmethod
    def set_run_fonts(run, font_config):
        """Passes the full font dictionary (ascii, cs, etc.)"""
        TextXml.set_run_fonts(run, font_config)

    @staticmethod
    def set_run_color(run_object, hex_color):
        """Can handle both High-level 'Run' object or Low-level 'element'."""
        if hasattr(run_object, '_element'):
            TextXml.set_run_color(run_object, hex_color) # Via standard flow
        else:
            TextXml.set_run_element_color(run_object, hex_color) # Direct element (Math)

    @staticmethod
    def set_run_underline_advanced(run, style='single', color=None):
        TextXml.set_underline_advanced(run, style, color)

    @staticmethod
    def set_run_shading(run, hex_color):
        TextXml.set_run_shading(run, hex_color)

    @staticmethod
    def set_run_border(run, **kwargs):
        """Args: size, val, color"""
        TextXml.set_run_border(run, **kwargs)

    @staticmethod
    def set_run_effects(run, spacing=0, shadow=None, glow=None, outline=None,reflection=None, gradient=None):
        """
        Public API Interface: Now supports Glow and Outline.
        Delegates call to the underlying TextXml factory.
        """
        # सुनिश्चित करें कि यहाँ 5 पैरामीटर्स (run, spacing, shadow, glow, outline) पास हो रहे हैं
        from .text_xml import TextXml
        TextXml.set_run_effects(
            run, 
            spacing=spacing, 
            shadow=shadow, 
            glow=glow, 
            outline=outline,
            reflection=reflection,
            gradient=gradient 
        )

    @staticmethod
    def append_child(parent_container, child_element):
        """Generic append safe for Oxml elements."""
        if hasattr(parent_container, '_element'):
            parent_container._element.append(child_element)
        else:
            parent_container.append(child_element) # Direct xml node

    @staticmethod
    def set_paragraph_outline_level(paragraph, level):
        TextXml.set_outline_level(paragraph, level)


    @staticmethod
    def set_run_scaling(run, scale_val):
        from .text_xml import TextXml
        TextXml.set_run_scaling(run, scale_val)

    @staticmethod
    def set_run_position(run, pos_val):
        from .text_xml import TextXml
        TextXml.set_run_position(run, pos_val)


    @staticmethod
    def set_run_shading_advanced(run, shading_data):
        """
        पोस्टमैन (XmlBuilder) ने डेटा पकड़ा और शेफ (TextXml) को दिया।
        खुद कुछ बनाने की कोशिश नहीं करनी है।
        """
        from .text_xml import TextXml
        # शेफ को आर्डर दो, शेफ के पास ही 'rPr' बनाने का औजार है
        TextXml.set_run_shading_advanced(run, shading_data)
        
        
    # =========================================================================
    # 2. ▦ TABLE OPERATIONS (Cells, Borders, Widths)
    # =========================================================================

    @staticmethod
    def set_cell_shading(cell, hex_color):
        TableXml.set_cell_shading(cell, hex_color)

    @staticmethod
    def set_cell_borders(cell, borders_dict):
        """borders_dict: {'top': {...}, 'left': {...}}"""
        TableXml.set_cell_borders(cell, borders_dict)

    @staticmethod
    def set_cell_v_merge(cell, val):
        TableXml.set_v_merge(cell, val)

    @staticmethod
    def set_cell_grid_span(cell, span):
        TableXml.set_grid_span(cell, span)

    @staticmethod
    def set_cell_margins(cell, margins_dict):
        TableXml.set_cell_margins(cell, margins_dict)

    @staticmethod
    def set_cell_valign(cell, align):
        """Top / Center / Bottom"""
        TableXml.set_vertical_alignment(cell, align)

    @staticmethod
    def set_cell_text_direction(cell, direction):
        """Rotates text inside cell."""
        TableXml.set_text_direction(cell, direction)

    @staticmethod
    def set_cell_width(cell, value, type_='dxa'):
        TableXml.set_cell_width(cell, value, type_)

    # --- Global Table Props ---
    @staticmethod
    def set_table_width_pct(table, pct_units):
        TableXml.set_table_width(table, pct_units, 'pct')

    @staticmethod
    def set_table_alignment(table, align_str):
        TableXml.set_table_alignment(table, align_str)

    @staticmethod
    def set_table_indent(table, twips):
        TableXml.set_table_indent(table, twips)

    @staticmethod
    def set_table_look(table, hex_val):
        TableXml.set_table_look(table, hex_val)

    @staticmethod
    def apply_invisible_layout_props(table, is_fixed=True):
        """ग्रिड/फ्लेक्स लेआउट के लिए टेबल को अदृश्य और फिक्स बनाता है।"""
        # 1. बॉर्डर हटाओ
        TableXml.set_table_borders_to_none(table)
        # 2. लेआउट फिक्स करो ताकि Percentages का सही पालन हो
        layout_mode = 'fixed' if is_fixed else 'autofit'
        TableXml.set_table_layout_preset(table, layout_mode)
        # 3. चौड़ाई 100% (5000 units) सेट करो (Standard Facade call)
        TableXml.set_table_width(table, 5000, 'pct')

    @staticmethod
    def set_grid_cell_margins(cell, top=120, bottom=120, left=180, right=180):
        """ग्रिड के भीतर मौजूद सेल के लिए संतुलित मार्जिन्स सेट करना।"""
        margins = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
        TableXml.set_cell_margins(cell, margins)


    def define_table_grid(table, widths_list):
        """Arguments: table (object), widths_list (list of ints in Twips)"""
        TableXml.define_table_grid(table, widths_list)

    def set_cell_no_wrap(cell, val_bool):
        """Wraps the internal TableXml method."""
        TableXml.set_cell_no_wrap(cell, val_bool)

    # =========================================================================
    # 3. 🖼️ DRAWING & MEDIA (The Assembly Line)
    # =========================================================================

    @classmethod
    def insert_image(cls, container, image_path, width_emu, height_emu, pos_config, image_id,style_data=None):
        """
        Complete Assembly for Images.
        Step 1: Get Relationship ID.
        Step 2: Generate <a:graphic>.
        Step 3: Wrap in Anchor/Inline based on pos_config.
        Step 4: Wrap in <w:drawing>.
        Step 5: Inject into <w:r> inside Container.
        """
        # 🔍 DEBUG
        debug_on = getattr(AppConfig, 'DEBUG_MEDIA', False)
        if debug_on:
            logger.debug(f"      🔨 [Builder] Inserting: Floating={pos_config.get('is_floating')}")

        
        # 1. Create Relationship (rId)
        part = container.part
        rId, _ = part.get_or_add_image(image_path)
        if debug_on: print(f"      🔗 rId Generated: {rId}")

        border_props = None
        if style_data and ('border' in style_data or any(k.startswith('border-') for k in style_data)):
            from kritidocx.basics.border_parser import BorderParser
            raw_border = style_data.get('border') or style_data.get('border-left')
            border_props = BorderParser.parse(raw_border)

        #  [DEBUG LOGGING START]
        if getattr(AppConfig, 'DEBUG_MEDIA', False):
            logger.debug(f"   🔨 [XML Builder] Handoff to DrawingXml:")
            logger.debug(f"      ➤ Border Props Parsed: {border_props}")
            if border_props:
                calc_emu = int(border_props.get('sz', 0) * 1587.5)
                logger.debug(f"      ➤ Expected Width in EMUs: {calc_emu}")
        #  [DEBUG LOGGING END]

        graphic = DrawingXml.create_image_graphic(rId, image_id, width_emu, height_emu, border_props=border_props)

        # 3. Determine Wrapper (Anchor vs Inline)
        is_floating = pos_config.get('is_floating', False)
        
        drawing_wrapper = None
        
        if is_floating:
            # Create Anchor with Positioning Math
            # Note: We need to augment pos_config with extents for the anchor frame
            anchor_config = pos_config.copy()
            anchor_config['width_emu'] = width_emu
            anchor_config['height_emu'] = height_emu
            
            # Step A: Build Base Anchor Tag
            anchor = DrawingXml.create_anchor(pos_config)
            
            # Step B: Apply Geometry & Position
            DrawingXml.apply_geometry(anchor, anchor_config)
            
            # Step C: Apply Wrapping
            DrawingXml.apply_wrapping(anchor, pos_config.get('wrap_type', 'square'))
            
            # Step D: Inject ID & Graphic
            docPr = cls.create_element('wp:docPr')
            uid = DrawingXml._get_next_shape_id()
            cls.create_attribute(docPr, 'id', str(uid))
            cls.create_attribute(docPr, 'name', f"Image {uid}")
            anchor.append(docPr)
            
            cNv = cls.create_element('wp:cNvGraphicFramePr')
            anchor.append(cNv)
            anchor.append(graphic)
            
            drawing_wrapper = anchor
            
            if debug_on: print("      ⚓ Mode: Anchor (Floating)")
            
        else:
            # Inline Wrapper
            inline = cls.create_element('wp:inline')
            for dist in ['distT', 'distB', 'distL', 'distR']:
                cls.create_attribute(inline, dist, "0")
            
            # 1. Extent (Size)
            extent = cls.create_element('wp:extent')
            cls.create_attribute(extent, 'cx', str(width_emu))
            cls.create_attribute(extent, 'cy', str(height_emu))
            inline.append(extent)
            
            # 🔥 [NUCLEAR FIX]: Effect Extent Expansion
            # यह Word को बताता है: "ड्राइंग, अपने बॉक्स से थोड़ा बाहर निकल सकती है।"
            # अगर हम इसे 0 रखते हैं, तो Word सख्त (Strict) होकर बॉर्डर के बाहरी किनारों को काट देता है।
            
            border_thickness = 0
            if border_props and border_props.get('sz'):
                # 1/8 pt to EMUs (approx 1587.5) -> थोड़ा बफर दें (1900 approx per point)
                border_thickness = int(border_props['sz'] * 1587.5)

            # डिफ़ॉल्ट 0, पर बॉर्डर हो तो मोटाई के बराबर एक्स्ट्रा जगह
            extent_val = str(border_thickness)
            
            effect = cls.create_element('wp:effectExtent')
            # Left, Top, Right, Bottom चारों तरफ मार्जिन बढ़ाएं
            cls.create_attribute(effect, 'l', extent_val)
            cls.create_attribute(effect, 't', extent_val)
            cls.create_attribute(effect, 'r', extent_val)
            cls.create_attribute(effect, 'b', extent_val)
            inline.append(effect)

            
            # 3. Doc Properties (ID & Name)
            docPr = cls.create_element('wp:docPr')
            uid = DrawingXml._get_next_shape_id()
            cls.create_attribute(docPr, 'id', str(uid))
            cls.create_attribute(docPr, 'name', f"Inline Image {uid}")
            inline.append(docPr)
            
            # 4. Content Properties (Locks etc)
            cNv = cls.create_element('wp:cNvGraphicFramePr')
            a_locks = cls.create_element('a:graphicFrameLocks') # 'a' namespace
            cls.create_attribute(a_locks, 'noChangeAspect', '1')
            cNv.append(a_locks)
            inline.append(cNv)
            
            # 5. The Graphic Itself
            inline.append(graphic)
            
            drawing_wrapper = inline

            if debug_on: print("      🧱 Mode: Inline (Standard)")
     
        # 4. Final Injection
        if drawing_wrapper is not None:
            # Create Run -> Drawing -> Wrapper
            run = container.add_run() if hasattr(container, 'add_run') else container
            drawing = cls.create_element('w:drawing')
            drawing.append(drawing_wrapper)
            
            if hasattr(run, '_r'): 
                run._r.append(drawing)
            else:
                run.append(drawing) # Direct XML element

            # 🔍 Verification
            if debug_on: print("      💉 [Builder] Appended <w:drawing> to Run.")
        else:
            if debug_on: print("      ⚠️ [Builder] ERROR: drawing_wrapper is None!")



    @classmethod
    def insert_textbox_shape(cls, container, pos_config, shape_style):
        """
        Assembly for Textbox.
        Returns: The internal text content node (<w:txbxContent>)
        so text logic can fill it.
        """
        # 1. Build Internal Graphic (WSP - Shape)
        # dims map
        dims = {
            'width_emu': pos_config.get('width', 1828800), # Default 2"
            'height_emu': pos_config.get('height', 914400) # Default 1"
        }
        graphic, txbxContent = DrawingXml.create_textbox_graphic(dims, shape_style)
        
        # 2. Build Anchor Wrapper
        pos_config_merged = {**pos_config, **dims} # Merge sizing for anchor extent
        anchor = DrawingXml.create_anchor(pos_config)
        DrawingXml.apply_geometry(anchor, pos_config_merged)
        
        # Determine wrapping based on floating boolean
        wrap = 'square' if pos_config.get('is_floating') else 'inline'
        # Force none wrap for shapes usually to sit on top freely unless text wrap specified
        if pos_config.get('wrap_type'): wrap = pos_config['wrap_type']
        
        DrawingXml.apply_wrapping(anchor, wrap)
        
        # 3. Add Identity
        docPr = cls.create_element('wp:docPr')
        uid = DrawingXml._get_next_shape_id()
        cls.create_attribute(docPr, 'id', str(uid))
        cls.create_attribute(docPr, 'name', f"Textbox {uid}")
        anchor.append(docPr)
        
        cNv = cls.create_element('wp:cNvGraphicFramePr')
        anchor.append(cNv)
        anchor.append(graphic) # Add Shape Graphic
        
        # 4. Inject into Container
        if hasattr(container, 'add_run'):
            # It's a paragraph object
            run = container.add_run()
            # Collapse run size (invisible anchor point)
            rPr = run._element.get_or_add_rPr()
            sz = cls.create_element('w:sz'); cls.create_attribute(sz, 'w:val', '1')
            rPr.append(sz)
            
            drawing = cls.create_element('w:drawing')
            drawing.append(anchor)
            run._r.append(drawing)
        else:
            # Raw append
            drawing = cls.create_element('w:drawing')
            drawing.append(anchor)
            container.append(drawing)
            
        return txbxContent # Pass this back to allow adding paragraphs inside!

    # =========================================================================
    # 4. 🎛️ FORM & FIELDS
    # =========================================================================

    @staticmethod
    def insert_sdt_checkbox(paragraph, checked, symbol_char, font_name, hex_color=None, font_size=24):
        FormXml.create_checkbox(paragraph, checked, symbol_char, font_name, hex_color, font_size)
    
    @staticmethod
    def insert_sdt_dropdown(paragraph, items, default_text):
        FormXml.create_dropdown(paragraph, items, default_text)

    @staticmethod
    def insert_sdt_text(paragraph, initial_text, is_placeholder, multiline=True): # Default True करें
        FormXml.create_text_input(paragraph, initial_text, is_placeholder, multiline)
        
    @staticmethod
    def insert_field_code(paragraph, code_text):
        """Example: PAGE, NUMPAGES"""
        FormXml.create_field_code(paragraph, code_text)

    # =========================================================================
    # 5. 📏 SECTION & LAYOUT
    # =========================================================================

    @staticmethod
    def set_page_size_xml(section, width, height, orientation='portrait'):
        # DEBUG: Facade Handoff Check
        logger.debug(f"   🏗️ [XmlBuilder] Handoff -> SectionXml: w={width}, h={height}, orient='{orientation}'")

        
        SectionXml.set_page_size(section, width, height, orientation)


    @staticmethod
    def set_section_margins_xml(section, margin_data):
        SectionXml.set_page_margins(section, margin_data)

    @staticmethod
    def set_section_columns(section, num, spacing_twips):
        SectionXml.set_columns(section, num, spacing_twips)

    @staticmethod
    def set_section_line_numbering(section, start, restart, count_by=1):
        # Pass the count_by value correctly to SectionXml
        SectionXml.set_line_numbering(section, start, count_by, restart)

    @staticmethod
    def set_section_borders(section, border_config):
        SectionXml.set_page_borders(section, border_config)

    @staticmethod
    def set_section_valign(section, align):
        SectionXml.set_vertical_alignment(section, align)

    # =========================================================================
    # 6. 🔢 NUMBERING SYSTEM
    # =========================================================================

    @staticmethod
    def register_abstract_list(part, abstract_id, levels):
        NumberingXml.register_abstract_list(part, abstract_id, levels)

    @staticmethod
    def register_list_instance(part, num_id, abstract_id):
        NumberingXml.register_list_instance(part, num_id, abstract_id)